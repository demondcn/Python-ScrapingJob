from __future__ import annotations

from pathlib import Path
import re
import unicodedata

from docx import Document

from .models import JobOffer
from .schemas import ResumeExperience, ResumeProfile, ResumeProject


TARGET_KEYWORDS: dict[str, list[str]] = {
    "devops_trainee": [
        "git",
        "github",
        "línea de comandos",
        "despliegue",
        "cloud",
        "automatización",
        "documentación",
        "scripts",
        "postgresql",
    ],
    "soporte_aplicaciones": [
        "soporte",
        "usuarios",
        "incidencias",
        "diagnóstico",
        "documentación",
        "inventario",
        "reportes",
        "aplicaciones web",
        "sql",
    ],
    "infraestructura_junior": [
        "soporte",
        "mantenimiento",
        "inventario",
        "equipos",
        "diagnóstico",
        "configuración",
        "incidencias",
        "línea de comandos",
    ],
    "cloud_support": [
        "cloud",
        "despliegue",
        "monitoreo",
        "logs",
        "soporte",
        "documentación",
        "postgresql",
    ],
    "qa_junior": [
        "documentación",
        "reportes",
        "validación",
        "incidencias",
        "aplicaciones web",
        "sql",
    ],
    "backend_junior": [
        "backend",
        "python",
        "java",
        "node.js",
        "express",
        "apis",
        "sql",
        "postgresql",
        "mysql",
        "git",
    ],
    "frontend_junior": [
        "frontend",
        "react",
        "next.js",
        "javascript",
        "typescript",
        "interfaces",
        "responsive",
        "vercel",
        "git",
    ],
    "fullstack_junior": [
        "fullstack",
        "react",
        "next.js",
        "javascript",
        "typescript",
        "node.js",
        "postgresql",
        "mysql",
        "git",
        "despliegue",
    ],
}


TARGET_SUMMARIES: dict[str, str] = {
    "devops_trainee": (
        "Tecnólogo en Desarrollo de Software con conocimientos en desarrollo web, bases de datos, Git, "
        "línea de comandos, despliegue de aplicaciones y soporte técnico. Experiencia práctica en proyectos "
        "con React, Next.js, PostgreSQL, Vercel y Neon, con interés en crecer en DevOps, automatización, "
        "infraestructura cloud y CI/CD."
    ),
    "soporte_aplicaciones": (
        "Tecnólogo en Desarrollo de Software con experiencia práctica en soporte técnico, desarrollo web, "
        "bases de datos y documentación. Con capacidad para diagnosticar incidencias, apoyar usuarios, revisar "
        "aplicaciones, generar reportes y colaborar en la mejora de procesos tecnológicos."
    ),
    "infraestructura_junior": (
        "Tecnólogo en Desarrollo de Software con experiencia en soporte operativo, trabajo con centros de "
        "cómputo, inventarios tecnológicos y mantenimiento básico de equipos. Interesado en crecer en "
        "infraestructura TI, administración de sistemas y automatización."
    ),
    "cloud_support": (
        "Tecnólogo en Desarrollo de Software con experiencia práctica en despliegue de aplicaciones, bases de "
        "datos, control de versiones y soporte técnico. Interesado en fortalecer su perfil en operación cloud, "
        "monitoreo, continuidad de servicios y administración de entornos."
    ),
    "qa_junior": (
        "Tecnólogo en Desarrollo de Software con experiencia en desarrollo web, documentación y seguimiento de "
        "procesos. Interesado en aportar en validación funcional, reporte de hallazgos, calidad de software y "
        "mejora continua."
    ),
    "backend_junior": (
        "Tecnólogo en Desarrollo de Software con experiencia práctica en desarrollo de aplicaciones, bases de "
        "datos y lógica de negocio. Con conocimientos en Python, Java, SQL y construcción de soluciones web "
        "orientadas al backend."
    ),
    "frontend_junior": (
        "TecnÃ³logo en Desarrollo de Software con experiencia prÃ¡ctica en desarrollo frontend, construcciÃ³n de "
        "interfaces, consumo de APIs y despliegue de proyectos web. Ha trabajado con React, Next.js, JavaScript, "
        "TypeScript y diseÃ±o responsivo en proyectos reales."
    ),
    "fullstack_junior": (
        "Tecnólogo en Desarrollo de Software con experiencia práctica en desarrollo web full stack, interfaces "
        "modernas, bases de datos y despliegue de aplicaciones. Ha trabajado con React, Next.js, JavaScript, "
        "TypeScript, PostgreSQL, Vercel y Neon en proyectos reales."
    ),
}


STOPWORDS = {
    "a",
    "al",
    "and",
    "con",
    "de",
    "del",
    "el",
    "en",
    "for",
    "in",
    "la",
    "las",
    "los",
    "of",
    "o",
    "para",
    "por",
    "pro",
    "se",
    "software",
    "the",
    "to",
    "un",
    "una",
    "web",
    "y",
}
VALID_SHORT_TOKENS = {"sql", "c#", "c++"}
TRAILING_CONNECTORS = {"y", "con", "de", "en", "para", "por"}
BLOCKED_DOC_FRAGMENTS = (
    "mas informacion",
    "referencias",
    "disponibilidad",
    "hobbies",
    "intereses",
    "daniel pinto",
    "agudelo",
    "jesus pavon",
    "jesus blanco",
    "salitre magico",
    "vehicul",
    "parque",
    "raw_text",
)
EXPERIENCE_PRIORITY_BY_TARGET: dict[str, list[str]] = {
    "soporte_aplicaciones": [
        "universidad de cundinamarca|auxiliar de oficina ii",
        "universidad de cundinamarca|tecnico en soporte de equipos de computo",
        "emprex360|desarrollador full stack web",
        "chocontano restaurante|desarrollador frontend y mobile",
        "charles barber|desarrollador android",
        "kepri holistica|desarrollador frontend web",
    ],
    "infraestructura_junior": [
        "universidad de cundinamarca|tecnico en soporte de equipos de computo",
        "universidad de cundinamarca|auxiliar de oficina ii",
        "emprex360|desarrollador full stack web",
        "chocontano restaurante|desarrollador frontend y mobile",
        "charles barber|desarrollador android",
        "kepri holistica|desarrollador frontend web",
    ],
    "devops_trainee": [
        "emprex360|desarrollador full stack web",
        "kepri holistica|desarrollador frontend web",
        "universidad de cundinamarca|auxiliar de oficina ii",
        "chocontano restaurante|desarrollador frontend y mobile",
        "charles barber|desarrollador android",
    ],
    "frontend_junior": [
        "kepri holistica|desarrollador frontend web",
        "chocontano restaurante|desarrollador frontend y mobile",
        "emprex360|desarrollador full stack web",
        "charles barber|desarrollador android",
        "universidad de cundinamarca|auxiliar de oficina ii",
    ],
}
CANONICAL_EXPERIENCE_BULLETS: dict[str, list[str]] = {
    "emprex360|desarrollador full stack web": [
        "Desarrollo de plataforma web para análisis de datos y diagnóstico empresarial.",
        "Implementación de dashboards interactivos y carga de datos en tiempo real con React, Next.js y PostgreSQL.",
        "Despliegue en Vercel con base de datos alojada en Neon.",
    ],
    "chocontano restaurante|desarrollador frontend y mobile": [
        "Desarrollo de aplicación móvil con React Native y sistema E-Commerce web.",
        "Implementación de pedidos en línea, inventario y panel administrativo.",
    ],
    "charles barber|desarrollador android": [
        "Desarrollo de aplicación Android en Java para control de inventario y ventas.",
        "Implementación de módulo de tienda en línea y gestión administrativa.",
    ],
    "universidad de cundinamarca|auxiliar de oficina ii": [
        "Apoyo operativo y soporte a usuarios en centros de cómputo.",
        "Gestión e inventario de recursos tecnológicos en laboratorios.",
        "Automatización de procesos de apagado de equipos para optimizar la operación.",
    ],
    "universidad de cundinamarca|tecnico en soporte de equipos de computo": [
        "Mantenimiento preventivo y correctivo de equipos de cómputo.",
        "Documentación de historial de mantenimiento, actualizaciones y configuraciones.",
        "Apoyo en diagnóstico de fallos de hardware y software.",
    ],
    "kepri holistica|desarrollador frontend web": [
        "Desarrollo de página web E-Commerce con Next.js.",
        "Implementación de catálogo de servicios, información institucional y diseño responsivo.",
    ],
}
SPANISH_REPLACEMENTS = [
    ("en linea", "en línea"),
    ("En linea", "En línea"),
    ("linea de comandos", "línea de comandos"),
    ("Linea de comandos", "Línea de comandos"),
    ("Tecnologo", "Tecnólogo"),
    ("tecnologo", "tecnólogo"),
    ("Tecnologia", "Tecnología"),
    ("tecnologia", "tecnología"),
    ("practica", "práctica"),
    ("Practica", "Práctica"),
    ("practicas", "prácticas"),
    ("Practicas", "Prácticas"),
    ("tecnico", "técnico"),
    ("Tecnico", "Técnico"),
    ("tecnicos", "técnicos"),
    ("Tecnicos", "Técnicos"),
    ("documentacion", "documentación"),
    ("Documentacion", "Documentación"),
    ("tecnologicos", "tecnológicos"),
    ("Tecnologicos", "Tecnológicos"),
    ("diagnostico", "diagnóstico"),
    ("Diagnostico", "Diagnóstico"),
    ("Educacion", "Educación"),
    ("educacion", "educación"),
    ("operacion", "operación"),
    ("Operacion", "Operación"),
    ("Linea", "Línea"),
    ("linea", "línea"),
    ("computo", "cómputo"),
    ("Computo", "Cómputo"),
    ("analisis", "análisis"),
    ("Analisis", "Análisis"),
    ("tecnico", "técnico"),
    ("Tecnico", "Técnico"),
    ("tecnicas", "técnicas"),
    ("Tecnicas", "Técnicas"),
    ("tecnica", "técnica"),
    ("Tecnica", "Técnica"),
    ("introduccion", "introducción"),
    ("Introduccion", "Introducción"),
    ("gestion", "gestión"),
    ("Gestion", "Gestión"),
    ("configuracion", "configuración"),
    ("Configuracion", "Configuración"),
    ("resolucion", "resolución"),
    ("Resolucion", "Resolución"),
    ("losequipos", "los equipos"),
    ("Losequipos", "Los equipos"),
    ("ingles", "inglés"),
    ("Ingles", "Inglés"),
    ("movil", "móvil"),
    ("Movil", "Móvil"),
    ("Autenticacion", "Autenticación"),
    ("autenticacion", "autenticación"),
    ("logica", "lógica"),
    ("Logica", "Lógica"),
]
KNOWN_TECH_PATTERNS: dict[str, tuple[str, ...]] = {
    "PostgreSQL": (r"\bpostgresql\b", r"\bpostgres\b"),
    "MySQL": (r"\bmysql\b",),
    "SQL": (r"\bsql\b",),
    "JavaScript": (r"\bjavascript\b",),
    "TypeScript": (r"\btypescript\b",),
    "React": (r"\breact\b",),
    "Next.js": (r"\bnext\.?js\b",),
    "Node.js": (r"\bnode\.?js\b",),
    "Express": (r"\bexpress\b",),
    "Python": (r"\bpython\b",),
    "Java": (r"\bjava\b",),
    "Git": (r"\bgit\b",),
    "GitHub": (r"\bgithub\b",),
    "Linux": (r"\blinux\b",),
    "Docker": (r"\bdocker\b",),
    "Vercel": (r"\bvercel\b",),
    "Neon": (r"\bneon\b",),
    "PHP": (r"\bphp\b",),
    "MongoDB": (r"\bmongodb\b",),
    "Redis": (r"\bredis\b",),
}


def get_available_targets() -> list[str]:
    return sorted(TARGET_KEYWORDS)


def build_ats_resume(
    profile: ResumeProfile,
    target: str,
    output_path: Path,
    job_offer: JobOffer | None = None,
) -> Path:
    if target not in TARGET_KEYWORDS:
        raise ValueError(f"Target no soportado: {target}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(_polish_text(profile.full_name) or "CV ATS", level=0)

    contact_line = _build_contact_line(profile)
    if contact_line:
        document.add_paragraph(contact_line)

    document.add_heading("Perfil Profesional", level=1)
    document.add_paragraph(build_target_summary(profile, target, job_offer))

    skill_sections = build_skill_categories(profile, target, job_offer)
    if skill_sections:
        document.add_heading("Habilidades Técnicas", level=1)
        for category, skills in skill_sections:
            if skills:
                document.add_paragraph(f"{category}: {', '.join(skills)}")

    experiences = select_relevant_experiences(profile, target, job_offer)
    projects = select_relevant_projects(profile, target, job_offer, experiences)
    if experiences or projects:
        document.add_heading("Experiencia Profesional y Proyectos", level=1)
        for experience in experiences:
            document.add_paragraph(_format_experience_header(experience))
            for bullet in experience.bullets[:3]:
                cleaned_bullet = _trim_bullet(bullet)
                if cleaned_bullet:
                    document.add_paragraph(cleaned_bullet, style="List Bullet")
        for project in projects:
            document.add_paragraph(_format_project_header(project))
            for bullet in project.bullets[:3]:
                cleaned_bullet = _trim_bullet(bullet)
                if cleaned_bullet:
                    document.add_paragraph(cleaned_bullet, style="List Bullet")
            if project.url:
                document.add_paragraph(_clean_text(project.url))

    education_items = _filter_non_relevant(profile.education)
    if education_items:
        document.add_heading("Educación", level=1)
        for item in education_items:
            document.add_paragraph(item, style="List Bullet")

    certification_items = _filter_non_relevant(profile.certifications)
    if certification_items:
        document.add_heading("Certificaciones y Formación Complementaria", level=1)
        for item in certification_items:
            document.add_paragraph(item, style="List Bullet")

    language_items = _filter_non_relevant(profile.languages)
    if language_items:
        document.add_heading("Idiomas", level=1)
        for item in language_items:
            document.add_paragraph(item, style="List Bullet")

    document.save(output_path)
    return output_path


def build_target_summary(profile: ResumeProfile, target: str, job_offer: JobOffer | None = None) -> str:
    base = TARGET_SUMMARIES[target]
    featured_skills = rank_skills(profile, target, job_offer)[:4]
    details: list[str] = []
    if featured_skills:
        details.append(f"Capacidades relevantes: {', '.join(featured_skills)}.")
    if job_offer and job_offer.title:
        details.append(f"Adaptado para una vacante de {job_offer.title}.")
    return _polish_text(" ".join([base, *details]).strip())


def rank_skills(profile: ResumeProfile, target: str, job_offer: JobOffer | None = None) -> list[str]:
    ranked: list[str] = []
    for _, skills in build_skill_categories(profile, target, job_offer):
        ranked.extend(skills)
    return _dedupe_keep_order(ranked)


def build_skill_categories(
    profile: ResumeProfile,
    target: str,
    job_offer: JobOffer | None = None,
) -> list[tuple[str, list[str]]]:
    context = _build_profile_context(profile)
    detected_tech = set(_detect_known_technologies_from_text(context))
    signals = _collect_signals(profile, context, detected_tech)
    sections = _build_target_sections(target, signals, detected_tech)
    if job_offer is not None:
        offer_keywords = _extract_offer_keywords(job_offer)
        sections = [(category, _sort_skills_by_offer(skills, offer_keywords)) for category, skills in sections]
    return [(category, _dedupe_keep_order(skills)) for category, skills in sections if skills]


def select_relevant_experiences(
    profile: ResumeProfile,
    target: str,
    job_offer: JobOffer | None = None,
) -> list[ResumeExperience]:
    target_keywords = TARGET_KEYWORDS[target]
    job_keywords = _extract_offer_keywords(job_offer)
    priority_list = EXPERIENCE_PRIORITY_BY_TARGET.get(target, [])
    priority_index = {value: index for index, value in enumerate(priority_list)}
    ranked = sorted(
        profile.experiences,
        key=lambda item: (
            priority_index.get(_experience_identity(item), len(priority_index) + 100),
            -_experience_score(item, target, target_keywords, job_keywords),
            _normalize_phrase(item.company),
            _normalize_phrase(item.role),
        ),
    )
    project_map = {_normalize_phrase(project.name): project for project in profile.projects}
    results: list[ResumeExperience] = []
    seen_roles: set[str] = set()
    for item in ranked:
        if not _is_allowed_experience(item):
            continue
        identity = _experience_identity(item)
        if identity in seen_roles:
            continue
        canonical_bullets = _get_canonical_experience_bullets(item)
        bullets = _filter_relevant_bullets(item.bullets, target_keywords, job_keywords)
        bullets = _augment_support_bullets(profile, item, bullets, target)
        related_project = project_map.get(_normalize_phrase(item.company))
        if canonical_bullets is None and related_project and related_project.url:
            project_bullet = f"Proyecto en línea: {_clean_text(related_project.url)}"
            if project_bullet not in bullets:
                bullets.append(project_bullet)
        cleaned_bullets = _prepare_experience_bullets(item, canonical_bullets, bullets)
        results.append(
            ResumeExperience(
                role=_polish_text(item.role),
                company=_polish_text(item.company),
                location=_polish_text(item.location),
                start_date=_polish_text(item.start_date),
                end_date=_polish_text(item.end_date),
                bullets=_dedupe_keep_order(cleaned_bullets),
            )
        )
        seen_roles.add(identity)
        if len(results) >= _experience_limit(target):
            break
    return results


def select_relevant_projects(
    profile: ResumeProfile,
    target: str,
    job_offer: JobOffer | None,
    experiences: list[ResumeExperience],
) -> list[ResumeProject]:
    used_keys = {_normalize_phrase(item.company) for item in experiences}
    target_keywords = TARGET_KEYWORDS[target]
    job_keywords = _extract_offer_keywords(job_offer)
    ranked = sorted(
        profile.projects,
        key=lambda item: _text_score(" ".join([item.name, item.role, *item.technologies, *item.bullets]), target_keywords, job_keywords),
        reverse=True,
    )
    results: list[ResumeProject] = []
    for item in ranked:
        if not _is_allowed_project(item):
            continue
        if _normalize_phrase(item.name) in used_keys:
            continue
        bullets = _filter_relevant_bullets(item.bullets, target_keywords, job_keywords)
        technologies = [tech for tech in _normalize_tech_list(item.technologies) if tech not in STOPWORDS]
        cleaned_bullets = [bullet for bullet in (_trim_bullet(bullet) for bullet in bullets) if bullet][:3]
        results.append(
            ResumeProject(
                name=_polish_text(item.name),
                role=_polish_text(item.role),
                technologies=_dedupe_keep_order(technologies[:8]),
                bullets=_dedupe_keep_order(cleaned_bullets),
                url=_clean_text(item.url),
            )
        )
    return results[:2]


def build_ats_filename(profile: ResumeProfile, target: str, job_offer: JobOffer | None = None) -> str:
    base_name = _safe_name(_clean_text(profile.full_name) or "CV_ATS")
    if job_offer is not None:
        title = _safe_name(_clean_text(job_offer.title) or target)
        return f"CV_ATS_{base_name}_Oferta_{job_offer.id}_{title}.docx"
    return f"CV_ATS_{base_name}_{_safe_name(target)}.docx"


def _build_contact_line(profile: ResumeProfile) -> str:
    parts = [
        _polish_text(profile.location),
        _clean_text(profile.phone),
        _clean_text(profile.email),
        _clean_text(profile.linkedin),
        _clean_text(profile.github),
        _clean_text(profile.portfolio),
    ]
    return " | ".join(part for part in parts if part and not _contains_blocked_doc_fragment(part))


def _experience_identity(experience: ResumeExperience) -> str:
    return "|".join([_normalize_phrase(experience.company), _normalize_phrase(experience.role)])


def _experience_limit(target: str) -> int:
    priority_list = EXPERIENCE_PRIORITY_BY_TARGET.get(target)
    if priority_list:
        return len(priority_list)
    return 5


def _get_canonical_experience_bullets(experience: ResumeExperience) -> list[str] | None:
    bullets = CANONICAL_EXPERIENCE_BULLETS.get(_experience_identity(experience))
    if bullets is None:
        return None
    return [_polish_text(bullet) for bullet in bullets]


def _prepare_experience_bullets(
    experience: ResumeExperience,
    canonical_bullets: list[str] | None,
    bullets: list[str],
) -> list[str]:
    if canonical_bullets is not None:
        return canonical_bullets[:3]

    expanded: list[str] = []
    for bullet in bullets:
        expanded.extend(_split_bullet_ideas(bullet))
    cleaned_bullets = [bullet for bullet in (_trim_bullet(bullet) for bullet in expanded) if bullet]
    return _dedupe_semantic_bullets(cleaned_bullets)[:3]


def _should_render_company_first(experience: ResumeExperience) -> bool:
    normalized_company = _normalize_phrase(experience.company)
    return normalized_company in {"emprex360", "chocontano restaurante", "charles barber", "kepri holistica"}


def _format_experience_header(experience: ResumeExperience) -> str:
    if _should_render_company_first(experience):
        lead = f"{experience.company} - {experience.role}".strip(" -")
    else:
        lead = f"{experience.role} - {experience.company}".strip(" -")
    segments = [lead]
    if experience.location:
        segments.append(experience.location)
    dates = " - ".join(part for part in [experience.start_date, experience.end_date] if part)
    if dates:
        segments.append(dates)
    return " | ".join(segment for segment in segments if segment)


def _format_project_header(project: ResumeProject) -> str:
    technologies = f" ({', '.join(project.technologies)})" if project.technologies else ""
    role = f" - {project.role}" if project.role else ""
    return f"{project.name}{role}{technologies}"


def _collect_signals(profile: ResumeProfile, context: str, detected_tech: set[str]) -> dict[str, object]:
    has_support_roles = any(
        _contains_any(
            _normalize_phrase(" ".join([item.role, item.company])),
            ("auxiliar de oficina", "tecnico en soporte", "técnico en soporte", "soporte", "centros de computo", "centros de cómputo"),
        )
        for item in profile.experiences
    )
    return {
        "_context": context,
        "support": has_support_roles or _contains_any(context, ("soporte tecnico", "soporte administrativo", "soporte basico", "soporte básico")),
        "support_users": _contains_any(context, ("usuarios", "uso de herramientas informaticas", "uso de equipos")),
        "diagnosis": _contains_any(context, ("diagnostico", "diagnóstico", "fallos", "incidencias", "problemas")),
        "documentation": _contains_any(context, ("documentacion", "documentación", "hojas de vida detalladas")),
        "reports": _contains_any(context, ("reportes", "reporte", "planillas")),
        "inventory": _contains_any(context, ("inventario", "inventarios", "recursos tecnologicos", "recursos tecnológicos")),
        "maintenance": _contains_any(context, ("mantenimiento preventivo", "mantenimiento correctivo", "hardware", "software", "historial de mantenimiento")),
        "equipment_config": _contains_any(context, ("configuracion", "configuración", "equipos de computo", "equipos de cómputo", "actualizaciones")),
        "networking": has_support_roles or _contains_any(context, ("conectividad", "laboratorios", "centros de computo", "centros de cómputo")),
        "command_line": _contains_any(context, ("linea de comandos", "línea de comandos", "terminal", "pip")),
        "deployment": _contains_any(context, ("despliegue", "vercel", "neon")),
        "cloud": _contains_any(context, ("cloud", "aws", "azure")) or {"Vercel", "Neon"} <= detected_tech,
        "automation": _contains_any(context, ("automatizacion", "automatización", "apagado de equipos", "scripts")),
        "problem_solving": _contains_any(context, ("resolucion de problemas", "resolución de problemas", "pensamiento analitico", "pensamiento analítico", "diagnostico")),
        "web_apps": _contains_any(context, ("aplicaciones web", "plataforma web", "e-commerce", "dashboards", "panel administrativo")),
        "env_vars": _contains_any(context, ("variables de entorno",)) or {"Vercel", "Neon"} <= detected_tech,
    }


def _build_target_sections(
    target: str,
    signals: dict[str, object],
    detected_tech: set[str],
) -> list[tuple[str, list[str]]]:
    def has_tech(name: str) -> bool:
        return name in detected_tech

    if target == "devops_trainee":
        return [
            (
                "DevOps básico",
                _skills_if(
                    ("Git", has_tech("Git")),
                    ("GitHub", has_tech("GitHub")),
                    ("Línea de comandos", signals["command_line"]),
                    ("Despliegue de aplicaciones", signals["deployment"]),
                    ("Variables de entorno", signals["env_vars"]),
                ),
            ),
            (
                "Cloud y despliegue",
                _skills_if(
                    ("Vercel", has_tech("Vercel")),
                    ("Neon", has_tech("Neon")),
                    ("Fundamentos de cloud", signals["cloud"]),
                ),
            ),
            (
                "Bases de datos",
                _skills_if(
                    ("SQL", has_tech("SQL")),
                    ("PostgreSQL", has_tech("PostgreSQL")),
                    ("MySQL", has_tech("MySQL")),
                ),
            ),
            (
                "Desarrollo",
                _skills_if(
                    ("JavaScript", has_tech("JavaScript")),
                    ("TypeScript", has_tech("TypeScript")),
                    ("React", has_tech("React")),
                    ("Next.js", has_tech("Next.js")),
                    ("Python", has_tech("Python")),
                ),
            ),
            (
                "Automatización",
                _skills_if(
                    ("Scripts", signals["automation"] or has_tech("Python")),
                    ("Documentación técnica", signals["documentation"]),
                    ("Resolución de problemas", signals["problem_solving"]),
                ),
            ),
        ]
    if target == "infraestructura_junior":
        return [
            (
                "Soporte e infraestructura",
                _skills_if(
                    ("Soporte técnico", signals["support"]),
                    ("Soporte a usuarios", signals["support_users"]),
                    ("Mantenimiento preventivo/correctivo", signals["maintenance"]),
                    ("Inventario tecnológico", signals["inventory"]),
                ),
            ),
            (
                "Sistemas y operación",
                _skills_if(
                    ("Línea de comandos", signals["command_line"]),
                    ("Configuración de equipos", signals["equipment_config"]),
                    ("Diagnóstico de fallos", signals["diagnosis"]),
                ),
            ),
            (
                "Redes básicas",
                _skills_if(
                    ("Conectividad", signals["networking"]),
                    ("Configuración básica", signals["equipment_config"] or signals["networking"]),
                    ("Resolución de incidencias", signals["diagnosis"] or signals["support"]),
                ),
            ),
            (
                "Herramientas",
                _skills_if(
                    ("Git", has_tech("Git")),
                    ("GitHub", has_tech("GitHub")),
                    ("Documentación técnica", signals["documentation"]),
                ),
            ),
        ]
    if target == "soporte_aplicaciones":
        return [
            (
                "Soporte y aplicaciones",
                _skills_if(
                    ("Soporte técnico", signals["support"]),
                    ("Soporte a usuarios", signals["support_users"]),
                    ("Diagnóstico de incidencias", signals["diagnosis"]),
                    ("Documentación técnica", signals["documentation"]),
                    ("Reportes", signals["reports"]),
                    ("Inventario tecnológico", signals["inventory"]),
                    ("Mantenimiento de equipos", signals["maintenance"]),
                ),
            ),
            (
                "Bases de datos",
                _skills_if(
                    ("SQL", has_tech("SQL")),
                    ("PostgreSQL", has_tech("PostgreSQL")),
                    ("MySQL", has_tech("MySQL")),
                ),
            ),
            (
                "Desarrollo web",
                _skills_if(
                    ("JavaScript", has_tech("JavaScript")),
                    ("React", has_tech("React")),
                    ("Next.js", has_tech("Next.js")),
                    ("TypeScript", has_tech("TypeScript")),
                    ("Aplicaciones web", signals["web_apps"]),
                ),
            ),
            (
                "Herramientas",
                _skills_if(
                    ("Git", has_tech("Git")),
                    ("GitHub", has_tech("GitHub")),
                    ("Vercel", has_tech("Vercel")),
                    ("Neon", has_tech("Neon")),
                ),
            ),
            (
                "Sistemas y operación",
                _skills_if(
                    ("Línea de comandos", signals["command_line"]),
                    ("Despliegue de aplicaciones", signals["deployment"]),
                ),
            ),
        ]
    if target == "cloud_support":
        return [
            (
                "Soporte y continuidad",
                _skills_if(
                    ("Soporte técnico", signals["support"]),
                    ("Documentación técnica", signals["documentation"]),
                    ("Reportes", signals["reports"]),
                    ("Resolución de incidencias", signals["diagnosis"]),
                ),
            ),
            (
                "Cloud y despliegue",
                _skills_if(
                    ("Vercel", has_tech("Vercel")),
                    ("Neon", has_tech("Neon")),
                    ("Fundamentos de cloud", signals["cloud"]),
                    ("Monitoreo", signals["reports"] or signals["documentation"]),
                    ("Logs", signals["deployment"]),
                ),
            ),
            (
                "Bases de datos",
                _skills_if(
                    ("SQL", has_tech("SQL")),
                    ("PostgreSQL", has_tech("PostgreSQL")),
                    ("MySQL", has_tech("MySQL")),
                ),
            ),
            (
                "Herramientas",
                _skills_if(
                    ("Git", has_tech("Git")),
                    ("GitHub", has_tech("GitHub")),
                    ("Linux", has_tech("Linux") or signals["command_line"]),
                    ("Línea de comandos", signals["command_line"]),
                ),
            ),
        ]
    if target == "qa_junior":
        return [
            (
                "Calidad y soporte",
                _skills_if(
                    ("Documentación técnica", signals["documentation"]),
                    ("Reportes", signals["reports"]),
                    ("Validación funcional", signals["web_apps"]),
                    ("Diagnóstico de incidencias", signals["diagnosis"]),
                ),
            ),
            (
                "Bases de datos",
                _skills_if(
                    ("SQL", has_tech("SQL")),
                    ("PostgreSQL", has_tech("PostgreSQL")),
                    ("MySQL", has_tech("MySQL")),
                ),
            ),
            (
                "Desarrollo web",
                _skills_if(
                    ("JavaScript", has_tech("JavaScript")),
                    ("React", has_tech("React")),
                    ("Next.js", has_tech("Next.js")),
                    ("Aplicaciones web", signals["web_apps"]),
                ),
            ),
            (
                "Herramientas",
                _skills_if(
                    ("Git", has_tech("Git")),
                    ("GitHub", has_tech("GitHub")),
                    ("Línea de comandos", signals["command_line"]),
                ),
            ),
        ]
    if target == "backend_junior":
        return [
            (
                "Desarrollo backend",
                _skills_if(
                    ("Python", has_tech("Python")),
                    ("Java", has_tech("Java")),
                    ("Node.js", has_tech("Node.js")),
                    ("Express", has_tech("Express")),
                    ("APIs REST", signals["web_apps"]),
                ),
            ),
            (
                "Bases de datos",
                _skills_if(
                    ("SQL", has_tech("SQL")),
                    ("PostgreSQL", has_tech("PostgreSQL")),
                    ("MySQL", has_tech("MySQL")),
                ),
            ),
            (
                "Herramientas",
                _skills_if(
                    ("Git", has_tech("Git")),
                    ("GitHub", has_tech("GitHub")),
                    ("Vercel", has_tech("Vercel")),
                    ("Neon", has_tech("Neon")),
                ),
            ),
            (
                "Operación y despliegue",
                _skills_if(
                    ("Línea de comandos", signals["command_line"]),
                    ("Despliegue de aplicaciones", signals["deployment"]),
                    ("Documentación técnica", signals["documentation"]),
                ),
            ),
        ]
    if target == "frontend_junior":
        return [
            (
                "Frontend y UI",
                _skills_if(
                    ("React", has_tech("React")),
                    ("Next.js", has_tech("Next.js")),
                    ("JavaScript", has_tech("JavaScript")),
                    ("TypeScript", has_tech("TypeScript")),
                    ("Interfaces", signals["web_apps"]),
                    ("DiseÃ±o responsivo", "responsive" in str(signals.get("_context", "")) or "responsivo" in str(signals.get("_context", ""))),
                ),
            ),
            (
                "IntegraciÃ³n web",
                _skills_if(
                    ("Consumo de APIs", signals["web_apps"]),
                    ("Vercel", has_tech("Vercel")),
                    ("Git", has_tech("Git")),
                    ("GitHub", has_tech("GitHub")),
                ),
            ),
            (
                "ConstrucciÃ³n de productos",
                _skills_if(
                    ("E-Commerce", "e-commerce" in str(signals.get("_context", ""))),
                    ("Componentes reutilizables", has_tech("React") or has_tech("Next.js")),
                    ("DocumentaciÃ³n tÃ©cnica", signals["documentation"]),
                ),
            ),
        ]
    return [
        (
            "Frontend y experiencia web",
            _skills_if(
                ("React", has_tech("React")),
                ("Next.js", has_tech("Next.js")),
                ("JavaScript", has_tech("JavaScript")),
                ("TypeScript", has_tech("TypeScript")),
            ),
        ),
        (
            "Backend y datos",
            _skills_if(
                ("Node.js", has_tech("Node.js")),
                ("Python", has_tech("Python")),
                ("SQL", has_tech("SQL")),
                ("PostgreSQL", has_tech("PostgreSQL")),
                ("MySQL", has_tech("MySQL")),
            ),
        ),
        (
            "Herramientas y despliegue",
            _skills_if(
                ("Git", has_tech("Git")),
                ("GitHub", has_tech("GitHub")),
                ("Vercel", has_tech("Vercel")),
                ("Neon", has_tech("Neon")),
            ),
        ),
        (
            "Construcción de productos",
            _skills_if(
                ("Dashboards", signals["reports"]),
                ("E-Commerce", "e-commerce" in str(signals.get("_context", ""))),
                ("Aplicaciones móviles", "movil" in str(signals.get("_context", "")) or "móvil" in str(signals.get("_context", ""))),
            ),
        ),
    ]


def _skills_if(*pairs: tuple[str, bool]) -> list[str]:
    return [_polish_text(skill) for skill, condition in pairs if condition]


def _sort_skills_by_offer(skills: list[str], offer_keywords: list[str]) -> list[str]:
    return sorted(
        skills,
        key=lambda item: (_text_score(item, [], offer_keywords), _normalize_phrase(item)),
        reverse=True,
    )


def _experience_score(
    experience: ResumeExperience,
    target: str,
    target_keywords: list[str],
    job_keywords: list[str],
) -> int:
    text = " ".join([experience.role, experience.company, experience.location, *experience.bullets])
    score = _text_score(text, target_keywords, job_keywords)
    normalized = _normalize_phrase(text)
    if target in {"soporte_aplicaciones", "infraestructura_junior", "cloud_support"} and _contains_any(
        normalized,
        ("soporte", "inventario", "mantenimiento", "centros de computo", "centros de cómputo"),
    ):
        score += 8
    if target == "devops_trainee" and _contains_any(normalized, ("despliegue", "vercel", "neon", "postgresql", "automatizacion", "automatización")):
        score += 8
    return score


def _extract_offer_keywords(job_offer: JobOffer | None) -> list[str]:
    if job_offer is None:
        return []
    text = _normalize_phrase(
        " ".join(
            [
                job_offer.title or "",
                job_offer.description or "",
                job_offer.requirements or "",
                job_offer.portal or "",
            ]
        )
    )
    raw_tokens = re.findall(r"[a-z0-9#+.]{1,}", text)
    cleaned_tokens: list[str] = []
    for token in raw_tokens:
        if token in STOPWORDS:
            continue
        if token == "r":
            continue
        if len(token) == 1:
            continue
        if len(token) < 3 and token not in VALID_SHORT_TOKENS:
            continue
        if token in {"web", "software"} or token.endswith(".vercel.app"):
            continue
        cleaned_tokens.append(token)
    cleaned_tokens.extend(_normalize_phrase(skill) for skill in _detect_known_technologies_from_text(text))
    return _dedupe_keep_order(cleaned_tokens)


def _split_bullet_ideas(value: str) -> list[str]:
    cleaned = _clean_text(value)
    cleaned = re.sub(r"^[•●â€¢\-\*\u2022\u25CF]+\s*", "", cleaned)
    if not cleaned:
        return []
    parts = re.split(r"(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚÑ])", cleaned)
    return [part.strip() for part in parts if part.strip()]


def _filter_relevant_bullets(bullets: list[str], target_keywords: list[str], job_keywords: list[str]) -> list[str]:
    filtered = sorted(
        bullets,
        key=lambda item: _text_score(item, target_keywords, job_keywords),
        reverse=True,
    )
    relevant = [bullet for bullet in filtered if _text_score(bullet, target_keywords, job_keywords) > 0]
    if not relevant:
        relevant = bullets[:]
    cleaned: list[str] = []
    for bullet in relevant:
        for fragment in _split_bullet_ideas(bullet):
            normalized = _trim_bullet(fragment)
            if normalized:
                cleaned.append(normalized)
    return _dedupe_semantic_bullets(cleaned)


def _augment_support_bullets(
    profile: ResumeProfile,
    experience: ResumeExperience,
    bullets: list[str],
    target: str,
) -> list[str]:
    normalized_role = _normalize_phrase(experience.role)
    if target not in {"soporte_aplicaciones", "infraestructura_junior", "cloud_support"}:
        return bullets
    if "auxiliar de oficina" not in normalized_role and "tecnico en soporte" not in normalized_role and "técnico en soporte" not in normalized_role:
        return bullets

    merged = [bullet for bullet in bullets if "Operador" not in bullet and "Salitre" not in bullet]
    raw_text = _normalize_phrase(profile.raw_text)
    candidate_pairs = [
        ("centros de computo", "Apoyo operativo y soporte administrativo en los centros de cómputo."),
        ("recursos tecnologicos", "Gestión y organización de recursos tecnológicos en laboratorios y centros de cómputo."),
        ("automatizacion de procesos de apagado", "Automatización de procesos de apagado de equipos para optimizar el consumo y la operación."),
        ("soporte basico a usuarios", "Soporte básico a usuarios en el uso de equipos y herramientas informáticas."),
        ("registro de inventarios", "Registro de inventarios, elaboración de reportes y seguimiento al estado de los equipos."),
        ("registro y control de inventarios", "Registro y control de inventarios, reportes y novedades del centro de cómputo."),
        ("mantenimiento preventivo y correctivo", "Mantenimiento preventivo y correctivo de hardware y software en equipos de cómputo."),
        ("hojas de vida detalladas", "Documentación del historial de mantenimiento, actualizaciones y configuraciones de equipos."),
    ]
    for fragment, sentence in candidate_pairs:
        if fragment in raw_text and not any(fragment in _normalize_phrase(existing) for existing in merged):
            merged.append(sentence)
    return _dedupe_keep_order([_trim_bullet(item) for item in merged if _trim_bullet(item)])


def _trim_bullet(value: str) -> str:
    if _contains_blocked_doc_fragment(value):
        return ""
    cleaned = re.sub(r"^[•●â€¢\-\*\u2022\u25CF]+\s*", "", value).strip()
    cleaned = _polish_text(cleaned).replace("Proyecto en linea:", "Proyecto en línea:")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    words = cleaned.split()
    if len(words) > 32:
        words = words[:32]
    while words and _normalize_phrase(words[-1].rstrip(".,;:")) in TRAILING_CONNECTORS:
        words.pop()
    trimmed = " ".join(words).strip().rstrip(",;:")
    if not trimmed:
        return ""
    normalized = _normalize_phrase(trimmed)
    if _contains_blocked_doc_fragment(normalized):
        return ""
    if normalized in STOPWORDS:
        return ""
    return trimmed


def _dedupe_semantic_bullets(items: list[str]) -> list[str]:
    result: list[str] = []
    signatures: list[set[str]] = []
    for item in items:
        signature = _bullet_signature(item)
        if not signature:
            continue
        if any(_is_similar_bullet(signature, existing) for existing in signatures):
            continue
        result.append(item)
        signatures.append(signature)
    return result


def _bullet_signature(value: str) -> set[str]:
    tokens = re.findall(r"[a-z0-9#+]+", _normalize_phrase(value))
    return {
        token
        for token in tokens
        if token not in STOPWORDS and len(token) >= 3 and token not in {"proyecto", "linea"}
    }


def _is_similar_bullet(current: set[str], existing: set[str]) -> bool:
    overlap = len(current & existing)
    if overlap == 0:
        return False
    if overlap >= min(len(current), len(existing)) and min(len(current), len(existing)) >= 4:
        return True
    return overlap >= 4 and (overlap / max(len(current), len(existing))) >= 0.75


def _filter_non_relevant(items: list[str]) -> list[str]:
    blocked = ("referencias", "hobbies", "pasatiempos", "familiares", "barras de habilidad")
    cleaned_items: list[str] = []
    for item in items:
        cleaned = _polish_text(item)
        normalized = _normalize_phrase(cleaned)
        if not cleaned:
            continue
        if _contains_blocked_doc_fragment(normalized):
            continue
        if any(token in normalized for token in blocked):
            continue
        if normalized in STOPWORDS:
            continue
        cleaned_items.append(cleaned)
    return _dedupe_keep_order(cleaned_items)


def _safe_name(value: str) -> str:
    parts = re.findall(r"[A-Za-z0-9]+", value)
    return "_".join(parts) if parts else "Documento"


def _build_profile_context(profile: ResumeProfile) -> str:
    fragments: list[str] = [
        profile.raw_text,
        profile.professional_summary,
        " ".join(profile.technical_skills),
        " ".join(profile.soft_skills),
        " ".join(profile.education),
        " ".join(profile.certifications),
        " ".join(profile.languages),
    ]
    for experience in profile.experiences:
        fragments.append(" ".join([experience.role, experience.company, experience.location, *experience.bullets]))
    for project in profile.projects:
        fragments.append(" ".join([project.name, project.role, *project.technologies, *project.bullets, project.url]))
    return _normalize_phrase(" ".join(_clean_text(fragment) for fragment in fragments if fragment))


def _detect_known_technologies_from_text(text: str) -> list[str]:
    normalized_text = _normalize_phrase(text)
    detected: list[str] = []
    for technology, patterns in KNOWN_TECH_PATTERNS.items():
        if any(re.search(pattern, normalized_text) for pattern in patterns):
            detected.append(technology)
    return _dedupe_keep_order(detected)


def _normalize_tech_list(items: list[str]) -> list[str]:
    normalized: list[str] = []
    for item in items:
        lowered = _normalize_phrase(item)
        if lowered in STOPWORDS or lowered in {"r", "software", "web"}:
            continue
        for technology, patterns in KNOWN_TECH_PATTERNS.items():
            if any(re.search(pattern, lowered) for pattern in patterns):
                normalized.append(technology)
                break
    return _dedupe_keep_order(normalized)


def _text_score(text: str, target_keywords: list[str], job_keywords: list[str]) -> int:
    normalized_text = _normalize_phrase(_clean_text(text))
    return sum(2 for keyword in target_keywords if _normalize_phrase(keyword) in normalized_text) + sum(
        1 for keyword in job_keywords if _normalize_phrase(keyword) in normalized_text
    )


def _contains_any(text: str, fragments: tuple[str, ...]) -> bool:
    return any(_normalize_phrase(fragment) in text for fragment in fragments)


def _clean_text(value: str) -> str:
    if not value:
        return ""
    repaired = value
    if "Ã" in repaired or "â" in repaired:
        try:
            repaired = repaired.encode("latin1").decode("utf-8")
        except UnicodeError:
            repaired = value
    repaired = repaired.replace("\u2011", "-").replace("\u2013", "-").replace("\u2014", "-")
    return re.sub(r"\s+", " ", repaired).strip()


def _polish_text(value: str) -> str:
    text = _clean_text(value)
    for source, target in SPANISH_REPLACEMENTS:
        text = re.sub(rf"\b{re.escape(source)}\b", target, text)
    return text.strip()


def _normalize_phrase(value: str) -> str:
    cleaned = _clean_text(value).lower()
    normalized = unicodedata.normalize("NFKD", cleaned)
    normalized = "".join(char for char in normalized if not unicodedata.combining(char))
    return normalized


def _contains_blocked_doc_fragment(value: str) -> bool:
    normalized = _normalize_phrase(value)
    return any(fragment in normalized for fragment in BLOCKED_DOC_FRAGMENTS)


def _is_allowed_experience(experience: ResumeExperience) -> bool:
    combined = " ".join(
        [
            experience.role,
            experience.company,
            experience.location,
            experience.start_date,
            experience.end_date,
            *experience.bullets,
        ]
    )
    normalized = _normalize_phrase(combined)
    if _contains_blocked_doc_fragment(normalized):
        return False
    return "operador" not in normalized


def _is_allowed_project(project: ResumeProject) -> bool:
    combined = " ".join([project.name, project.role, *project.technologies, *project.bullets, project.url])
    return not _contains_blocked_doc_fragment(combined)


def _dedupe_keep_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        normalized = re.sub(r"[^\w\s#+]+", "", _normalize_phrase(item))
        if not item or not normalized:
            continue
        if normalized in STOPWORDS:
            continue
        if normalized == "r":
            continue
        if len(normalized) == 1:
            continue
        if len(normalized) < 3 and normalized not in VALID_SHORT_TOKENS:
            continue
        if normalized in {"web", "software"} or normalized.endswith(".vercel.app"):
            continue
        if normalized not in seen:
            result.append(item)
            seen.add(normalized)
    return result
