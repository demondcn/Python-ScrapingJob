from __future__ import annotations

from pathlib import Path
import re

from docx import Document

from .models import CandidateProfile, GeneratedDocument, JobOffer


def _safe_slug(value: str) -> str:
    clean = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return clean or "cv"


def generate_cv(output_dir: Path, profile: CandidateProfile, offer: JobOffer, template_path: Path | None = None) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)

    if template_path and template_path.exists():
        document = Document(template_path)
    else:
        document = Document()

    document.add_heading(profile.full_name, level=0)
    document.add_paragraph(f"Email: {profile.email}")
    document.add_paragraph(f"Telefono: {profile.phone}")
    document.add_paragraph(f"Ciudad: {profile.city}")

    document.add_heading("Resumen", level=1)
    document.add_paragraph(_build_summary(profile, offer))

    document.add_heading("Habilidades", level=1)
    document.add_paragraph(profile.skills)

    document.add_heading("Proyectos", level=1)
    document.add_paragraph(profile.projects)

    document.add_heading("Educacion", level=1)
    document.add_paragraph(profile.education)

    file_name = f"{_safe_slug(offer.company)}-{_safe_slug(offer.title)}.docx"
    output_path = output_dir / file_name
    document.save(output_path)
    return output_path


def _build_summary(profile: CandidateProfile, offer: JobOffer) -> str:
    fragments = [
        profile.summary.strip(),
        f"Perfil orientado a la vacante de {offer.title}.",
    ]
    if offer.match_reason:
        first_reason = offer.match_reason.splitlines()[0]
        fragments.append(f"Coincidencia principal: {first_reason}.")
    return " ".join(fragment for fragment in fragments if fragment)


def register_generated_cv(session, offer: JobOffer, file_path: Path) -> GeneratedDocument:
    record = GeneratedDocument(job_offer_id=offer.id, doc_type="cv", file_path=str(file_path))
    session.add(record)
    session.commit()
    session.refresh(record)
    return record

