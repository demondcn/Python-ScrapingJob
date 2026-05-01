from __future__ import annotations

from pathlib import Path
import re
import unicodedata

from docx import Document
from pypdf import PdfReader

from .schemas import ResumeExperience, ResumeProfile, ResumeProject


EMAIL_RE = re.compile(r"[\w.\-+]+\s*@\s*[\w.\-]+\.\w+")
PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")
URL_RE = re.compile(r"(https?://\S+|(?:linkedin|github)\.com/\S+|(?:www\.)?\S+\.\w{2,}/\S*)", re.IGNORECASE)
BULLET_PREFIX_RE = re.compile(r"^[•●\-\*\u2022\u25CF]\s*")
MULTISPACE_RE = re.compile(r"[ \t]{2,}")
CONNECTOR_ENDINGS = {"y", "con", "de", "en", "para", "por"}
ROLE_HINTS = (
    "auxiliar",
    "técnico",
    "tecnico",
    "desarrollador",
    "operador",
    "analista",
    "soporte",
    "ingeniero",
)
KNOWN_INLINE_SKILLS = [
    "Python",
    "Java",
    "C++",
    "C#",
    "JavaScript",
    "TypeScript",
    "React",
    "React Native",
    "Next.js",
    "Node.js",
    "Express",
    "PHP",
    "SQL",
    "PostgreSQL",
    "MySQL",
    "Git",
    "GitHub",
    "Linux",
    "Docker",
    "Vercel",
    "Neon",
    "MongoDB",
    "Redis",
]
BLOCKED_PROFILE_FRAGMENTS = (
    "referencias",
    "mas informacion",
    "más información",
    "disponibilidad",
    "hobbies",
    "intereses",
    "daniel pinto",
    "agudelo davalos",
    "jesus pavon",
    "jesus blanco",
    "salitre magico",
    "vehiculos",
    "vehículos",
    "parque",
    "raw_text",
)
LANGUAGE_HINTS = ("español", "espanol", "inglés", "ingles", "frances", "francés", "nativo", "lectura técnica", "lectura tecnica", "a1", "a2", "b1", "b2", "c1", "c2")
CERTIFICATION_HINTS = (
    "curso",
    "platzi",
    "frontend developer",
    "git y github",
    "linea de comandos",
    "línea de comandos",
    "javascript desde cero",
    "python",
    "java se",
    "mysql",
    "mariadb",
    "kotlin",
    "bases de datos",
)
MONTH_ORDER = {
    "enero": 1,
    "febrero": 2,
    "marzo": 3,
    "abril": 4,
    "mayo": 5,
    "junio": 6,
    "julio": 7,
    "agosto": 8,
    "septiembre": 9,
    "octubre": 10,
    "noviembre": 11,
    "diciembre": 12,
}

SECTION_MAP = {
    "perfil profesional": "summary",
    "perfil": "summary",
    "resumen profesional": "summary",
    "habilidades y competencias": "skills",
    "habilidades tecnicas": "skills",
    "habilidades": "skills",
    "experiencia profesional": "experience",
    "experiencia": "experience",
    "proyectos": "projects",
    "educacion": "education",
    "formacion complementaria": "certifications",
    "certificaciones": "certifications",
    "idiomas": "languages",
    "referencias": "references",
}


def read_resume_file(path: Path) -> ResumeProfile:
    if not path.exists():
        raise FileNotFoundError(f"El archivo no existe: {path}")
    suffix = path.suffix.lower()
    if suffix == ".docx":
        raw_text = _extract_text_from_docx(path)
    elif suffix == ".pdf":
        raw_text = _extract_text_from_pdf(path)
    else:
        raise ValueError(f"Formato no soportado: {suffix}. Usa .docx o .pdf")
    cleaned = clean_resume_text(raw_text)
    profile = parse_resume_profile(cleaned, source_path=path)
    for related_path in _find_related_resume_files(path, profile):
        related_profile = _read_single_resume_profile(related_path)
        profile = _merge_profiles(profile, related_profile)
    return _sanitize_resume_profile(profile)


def _read_single_resume_profile(path: Path) -> ResumeProfile:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        raw_text = _extract_text_from_docx(path)
    elif suffix == ".pdf":
        raw_text = _extract_text_from_pdf(path)
    else:
        raise ValueError(f"Formato no soportado: {suffix}. Usa .docx o .pdf")
    return parse_resume_profile(clean_resume_text(raw_text), source_path=path)


def clean_resume_text(text: str) -> str:
    normalized = _repair_text_encoding(text)
    normalized = normalized.replace("\r", "\n").replace("\x00", " ")
    normalized = normalized.replace("\u2011", "-").replace("\u2013", "-").replace("\u2014", "-")
    normalized = normalized.replace("\u00a0", " ")
    normalized = re.sub(r"[^\S\n]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    lines = [_normalize_line(line) for line in normalized.splitlines()]
    merged_lines: list[str] = []
    for line in lines:
        if not line:
            if merged_lines and merged_lines[-1] != "":
                merged_lines.append("")
            continue
        if (
            merged_lines
            and merged_lines[-1]
            and not _is_heading_line(merged_lines[-1])
            and not _is_bullet_line(line)
            and _should_join_lines(merged_lines[-1], line)
        ):
            merged_lines[-1] = f"{merged_lines[-1]} {line}".strip()
        else:
            merged_lines.append(line)
    return "\n".join(merged_lines).strip()


def parse_resume_profile(text: str, source_path: Path | None = None) -> ResumeProfile:
    lines = [line.strip() for line in text.splitlines()]
    non_empty = [line for line in lines if line]
    profile = ResumeProfile(raw_text=text)
    if non_empty:
        profile.full_name = _resolve_full_name(non_empty[0], source_path)
    profile.email = _normalize_email(_search_first(EMAIL_RE, text))
    profile.phone = _search_first(PHONE_RE, text)
    urls = URL_RE.findall(text)
    profile.linkedin = _pick_url(urls, "linkedin.com")
    profile.github = _pick_url(urls, "github.com")
    profile.portfolio = _pick_portfolio(urls)
    profile.location = _extract_location(non_empty[1] if len(non_empty) > 1 else "")

    sections = _split_sections(lines)
    profile.professional_summary = _extract_summary(sections)
    technical_skills, soft_skills = _extract_skills(sections.get("skills", []))
    if not technical_skills:
        technical_skills = _extract_inline_skills(lines)
    profile.technical_skills = technical_skills
    profile.soft_skills = soft_skills
    profile.experiences = _extract_experiences(sections.get("experience", []))
    if not profile.experiences:
        profile.experiences = _extract_experiences_fallback(lines)
    profile.projects = _extract_projects(sections.get("projects", []), profile.experiences)
    if not profile.professional_summary:
        profile.professional_summary = _extract_summary_fallback(lines, profile.experiences)
    profile.education = _clean_section_items(sections.get("education", []))
    profile.certifications = _clean_section_items(sections.get("certifications", []))
    profile.languages = _clean_section_items(sections.get("languages", []))
    return profile


def _extract_text_from_docx(path: Path) -> str:
    document = Document(path)
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def _extract_text_from_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _normalize_line(line: str) -> str:
    line = MULTISPACE_RE.sub(" ", line.strip())
    line = re.sub(r"\s+\|\s+", " | ", line)
    return line


def _normalize_key(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    normalized = "".join(char for char in normalized if not unicodedata.combining(char))
    return normalized.lower().strip(" :")


def _is_heading_line(line: str) -> bool:
    normalized = _normalize_key(line)
    return normalized in SECTION_MAP or (line.isupper() and len(line.split()) <= 5)


def _is_bullet_line(line: str) -> bool:
    return bool(BULLET_PREFIX_RE.match(line))


def _should_join_lines(previous: str, current: str) -> bool:
    if previous.endswith((".", ":", "|")):
        return False
    lower_current = current.lower()
    if any(token in lower_current for token in ("email:", "tel:", "telefono:", "linkedin:", "github:")):
        return False
    if "|" in current:
        return False
    if current.endswith(":"):
        return False
    if current.startswith(("http://", "https://")):
        return False
    if _is_heading_line(current):
        return False
    return previous[:1].isalnum() and current[:1].isalnum()


def _split_sections(lines: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current = "header"
    sections[current] = []
    for line in lines:
        if not line:
            sections.setdefault(current, []).append("")
            continue
        normalized = _normalize_key(line)
        if normalized in SECTION_MAP:
            current = SECTION_MAP[normalized]
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return sections


def _search_first(pattern: re.Pattern[str], text: str) -> str:
    match = pattern.search(text)
    if not match:
        return ""
    return match.group(1) if match.lastindex else match.group(0)


def _pick_url(urls: list[str], token: str) -> str:
    for url in urls:
        if token in url.lower():
            return _normalize_url(url)
    return ""


def _pick_portfolio(urls: list[str]) -> str:
    for url in urls:
        lower = url.lower()
        if "linkedin.com" not in lower and "github.com" not in lower:
            return _normalize_url(url)
    return ""


def _normalize_url(url: str) -> str:
    cleaned = re.sub(r"\s+", "", url).strip().rstrip(".,);·")
    if cleaned.startswith("http://") or cleaned.startswith("https://"):
        return cleaned
    return f"https://{cleaned.lstrip('/')}"


def _extract_location(line: str) -> str:
    if "|" in line:
        return line.split("|", 1)[0].strip()
    return line.strip()


def _extract_summary(sections: dict[str, list[str]]) -> str:
    summary_lines = _clean_section_items(sections.get("summary", []))
    return " ".join(summary_lines).strip()


def _extract_skills(lines: list[str]) -> tuple[list[str], list[str]]:
    technical: list[str] = []
    soft: list[str] = []
    current_bucket = technical
    for line in lines:
        if not line:
            continue
        normalized = _normalize_key(line)
        if "habilidades blandas" in normalized:
            current_bucket = soft
            continue
        if normalized.endswith("lenguajes y tecnologias") or normalized.endswith("areas de desarrollo") or normalized.endswith("fundamentos tecnicos"):
            current_bucket = technical
            continue
        if ":" in line:
            _, value = line.split(":", 1)
            items = _split_skill_items(value)
        else:
            items = _split_skill_items(line)
        current_bucket.extend(item for item in items if item)
    return _dedupe_keep_order(technical), _dedupe_keep_order(soft)


def _extract_experiences(lines: list[str]) -> list[ResumeExperience]:
    entries: list[ResumeExperience] = []
    current: ResumeExperience | None = None
    for line in lines:
        if not line:
            continue
        if _is_bullet_line(line):
            bullet = BULLET_PREFIX_RE.sub("", line).strip()
            if current is not None and bullet:
                current.bullets.append(_compact_sentence(bullet))
            continue
        parsed_header = _parse_experience_header(line)
        if parsed_header:
            if current is not None:
                current.bullets = [_trim_sentence(bullet) for bullet in current.bullets if _trim_sentence(bullet)]
                entries.append(current)
            current = parsed_header
            continue
        if current is not None:
            if _looks_like_dates_line(line):
                current.location, current.start_date, current.end_date = _parse_location_and_dates(line)
            elif line.startswith("Proyecto en línea:") or line.lower().startswith("proyecto en linea:"):
                current.bullets.append(_compact_sentence(line))
            elif current.bullets and (_ends_with_connector(current.bullets[-1]) or len(line.split()) <= 3):
                current.bullets[-1] = _compact_sentence(f"{current.bullets[-1]} {line}")
            elif len(line.split()) > 2:
                current.bullets.append(_compact_sentence(line))
    if current is not None:
        current.bullets = [_trim_sentence(bullet) for bullet in current.bullets if _trim_sentence(bullet)]
        entries.append(current)
    return entries


def _extract_experiences_fallback(lines: list[str]) -> list[ResumeExperience]:
    entries: list[ResumeExperience] = []
    current: ResumeExperience | None = None
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if _looks_like_role_title(line):
            if current is not None:
                current.bullets = [_trim_sentence(bullet) for bullet in current.bullets if _trim_sentence(bullet)]
                if current.company or current.bullets:
                    entries.append(current)
            role = line
            company = ""
            location = ""
            start_date = ""
            end_date = ""
            extra_bullet = ""
            lookahead = _next_non_empty_line(lines, index + 1)
            if lookahead and ("|" in lookahead or _looks_like_dates_line(lookahead)):
                company, location, start_date, end_date, extra_bullet = _parse_company_dates_line(lookahead)
                index = _index_of_next_non_empty_line(lines, index + 1)
            current = ResumeExperience(
                role=role,
                company=company,
                location=location,
                start_date=start_date,
                end_date=end_date,
                bullets=[_compact_sentence(extra_bullet)] if extra_bullet else [],
            )
            index += 1
            continue
        if current is not None:
            if _looks_like_role_title(line):
                continue
            if _looks_like_stop_section(line):
                current.bullets = [_trim_sentence(bullet) for bullet in current.bullets if _trim_sentence(bullet)]
                if current.company or current.bullets:
                    entries.append(current)
                current = None
                index += 1
                continue
            if _looks_like_inline_skill_line(line):
                index += 1
                continue
            if current.bullets and (_ends_with_connector(current.bullets[-1]) or len(line.split()) <= 4):
                current.bullets[-1] = _compact_sentence(f"{current.bullets[-1]} {line}")
            elif len(line.split()) >= 3:
                current.bullets.append(_compact_sentence(line))
        index += 1
    if current is not None:
        current.bullets = [_trim_sentence(bullet) for bullet in current.bullets if _trim_sentence(bullet)]
        if current.company or current.bullets:
            entries.append(current)
    return entries


def _parse_experience_header(line: str) -> ResumeExperience | None:
    if " — " in line:
        company, role = [part.strip() for part in line.split(" — ", 1)]
        role, location, start_date, end_date = _split_role_dates(role)
        return ResumeExperience(role=role, company=company, location=location, start_date=start_date, end_date=end_date)
    if " - " in line and any(token in line.lower() for token in ("desarrollador", "analista", "soporte", "ingeniero", "auxiliar")):
        company, role = [part.strip() for part in line.split(" - ", 1)]
        role, location, start_date, end_date = _split_role_dates(role)
        return ResumeExperience(role=role, company=company, location=location, start_date=start_date, end_date=end_date)
    return None


def _looks_like_role_title(line: str) -> bool:
    normalized = _normalize_key(line)
    if len(line.split()) > 8:
        return False
    if any(hint in normalized for hint in ROLE_HINTS):
        return True
    return False


def _parse_company_dates_line(line: str) -> tuple[str, str, str, str, str]:
    company = ""
    location = ""
    start_date = ""
    end_date = ""
    extra_bullet = ""
    if "|" in line:
        left, right = [part.strip() for part in line.split("|", 1)]
        company = left.rstrip(".")
        if _looks_like_dates_line(right):
            date_match = re.search(
                r"((?:Enero|Febrero|Marzo|Abril|Mayo|Junio|Julio|Agosto|Septiembre|Octubre|Noviembre|Diciembre)\s+\d{4})\s*-\s*((?:Enero|Febrero|Marzo|Abril|Mayo|Junio|Julio|Agosto|Septiembre|Octubre|Noviembre|Diciembre)\s+\d{4})",
                right,
                re.IGNORECASE,
            )
            if date_match:
                start_date, end_date = date_match.group(1), date_match.group(2)
                extra_bullet = right.replace(date_match.group(0), "", 1).strip(" .-")
            else:
                _, start_date, end_date = _parse_location_and_dates(right)
        else:
            location = right
    else:
        company = line.strip().rstrip(".")
    return company, location, start_date, end_date, extra_bullet


def _next_non_empty_line(lines: list[str], start_index: int) -> str:
    for line in lines[start_index:]:
        stripped = line.strip()
        if stripped:
            return stripped
    return ""


def _index_of_next_non_empty_line(lines: list[str], start_index: int) -> int:
    for index in range(start_index, len(lines)):
        if lines[index].strip():
            return index
    return len(lines) - 1


def _looks_like_dates_line(line: str) -> bool:
    months = (
        "enero",
        "febrero",
        "marzo",
        "abril",
        "mayo",
        "junio",
        "julio",
        "agosto",
        "septiembre",
        "octubre",
        "noviembre",
        "diciembre",
    )
    lower = _normalize_key(line)
    return any(month in lower for month in months) and any(char.isdigit() for char in lower)


def _parse_location_and_dates(line: str) -> tuple[str, str, str]:
    if "|" in line:
        location, dates = [part.strip() for part in line.split("|", 1)]
    else:
        location, dates = "", line.strip()
    if "–" in dates:
        start, end = [part.strip() for part in dates.split("–", 1)]
    elif "-" in dates:
        start, end = [part.strip() for part in dates.split("-", 1)]
    else:
        start, end = dates, ""
    return location, start, end


def _split_role_dates(value: str) -> tuple[str, str, str, str]:
    location = ""
    role_text = value.strip()
    location_match = re.search(r"\(([^)]+)\)", role_text)
    if location_match:
        location = location_match.group(1).strip()
        role_text = role_text.replace(location_match.group(0), "").strip()
    date_match = re.search(
        r"((?:Enero|Febrero|Marzo|Abril|Mayo|Junio|Julio|Agosto|Septiembre|Octubre|Noviembre|Diciembre)\s+\d{4})\s*-\s*((?:Enero|Febrero|Marzo|Abril|Mayo|Junio|Julio|Agosto|Septiembre|Octubre|Noviembre|Diciembre)\s+\d{4})",
        role_text,
        re.IGNORECASE,
    )
    start_date = ""
    end_date = ""
    if date_match:
        start_date, end_date = date_match.group(1), date_match.group(2)
        role_text = role_text.replace(date_match.group(0), "").strip(" -")
    return role_text.strip(), location, start_date, end_date


def _extract_projects(lines: list[str], experiences: list[ResumeExperience]) -> list[ResumeProject]:
    projects: list[ResumeProject] = []
    for line in lines:
        if not line:
            continue
        if "http" in line.lower():
            projects.append(ResumeProject(name=line.split("http", 1)[0].strip(" :"), url="http" + line.split("http", 1)[1]))
    for experience in experiences:
        url = ""
        bullets: list[str] = []
        for bullet in experience.bullets:
            url_match = re.search(r"https?://\S+", bullet)
            if url_match:
                url = url_match.group(0).rstrip(".,)")
            else:
                bullets.append(bullet)
        if url or _looks_like_project_experience(experience):
            technologies = _extract_known_technologies(" ".join([experience.role, *experience.bullets]))
            projects.append(
                ResumeProject(
                    name=experience.company,
                    role=experience.role,
                    technologies=technologies,
                    bullets=bullets[:4],
                    url=url,
                )
            )
    return projects


def _extract_inline_skills(lines: list[str]) -> list[str]:
    joined = " ".join(lines)
    found: list[str] = []
    lower_joined = joined.lower()
    for skill in KNOWN_INLINE_SKILLS:
        if skill.lower() in lower_joined:
            found.append(skill)
    return _dedupe_keep_order(found)


def _looks_like_project_experience(experience: ResumeExperience) -> bool:
    text = " ".join([experience.company, experience.role, *experience.bullets]).lower()
    return any(token in text for token in ("react", "next.js", "dashboard", "e-commerce", "android", "aplicacion", "plataforma web"))


def _extract_known_technologies(text: str) -> list[str]:
    known = [
        "Python",
        "Java",
        "JavaScript",
        "TypeScript",
        "React",
        "Next.js",
        "React Native",
        "PostgreSQL",
        "MySQL",
        "MongoDB",
        "Redis",
        "Vercel",
        "Neon",
        "SQL",
        "Node.js",
        "Express",
    ]
    lower_text = text.lower()
    return [item for item in known if item.lower() in lower_text]


def _clean_section_items(lines: list[str]) -> list[str]:
    items: list[str] = []
    for line in lines:
        cleaned = BULLET_PREFIX_RE.sub("", line).strip()
        if not cleaned:
            continue
        if "referenc" in _normalize_key(cleaned):
            continue
        items.append(_trim_sentence(cleaned))
    return items


def _extract_summary_fallback(lines: list[str], experiences: list[ResumeExperience]) -> str:
    role_titles = {_normalize_key(item.role) for item in experiences}
    summary_candidates: list[str] = []
    for line in lines:
        normalized = _normalize_key(line)
        if normalized in role_titles:
            break
        if len(line.split()) >= 12 and "@" not in line and "http" not in line:
            summary_candidates.append(line.strip())
    return " ".join(summary_candidates[:2]).strip()


def _split_skill_items(value: str) -> list[str]:
    value = value.replace("(", ", ").replace(")", "")
    tokens = re.split(r"[•,;|]", value)
    return [_trim_sentence(token) for token in tokens if token.strip()]


def _trim_sentence(value: str, max_words: int | None = None) -> str:
    words = value.strip().split()
    if max_words and len(words) > max_words:
        words = words[:max_words]
    text = " ".join(words).strip()
    return _strip_trailing_connectors(text)


def _compact_sentence(value: str) -> str:
    return " ".join(value.strip().split()).strip()


def _repair_text_encoding(text: str) -> str:
    if "Ã" not in text and "â" not in text:
        return text
    try:
        repaired = text.encode("latin1").decode("utf-8")
    except UnicodeError:
        return text
    return repaired if repaired.count("Ã") <= text.count("Ã") else text


def _normalize_email(email: str) -> str:
    return re.sub(r"\s+", "", email)


def _resolve_full_name(first_line: str, source_path: Path | None) -> str:
    if not source_path:
        return _strip_contact_noise_from_name(first_line.strip())
    normalized = first_line.strip()
    normalized = _strip_contact_noise_from_name(normalized)
    if re.search(r"\b[A-Z]\.$", normalized):
        file_name = source_path.stem.replace("_", " ")
        file_name = re.sub(r"^(HV|CV)\s+", "", file_name, flags=re.IGNORECASE)
        file_name = re.sub(r"\s+", " ", file_name).strip()
        if len(file_name.split()) >= 4:
            return file_name
    return normalized


def _looks_like_inline_skill_line(line: str) -> bool:
    normalized = line.strip()
    if len(normalized.split()) > 8:
        return False
    upper_ratio = sum(1 for char in normalized if char.isupper()) / max(len(normalized), 1)
    return upper_ratio > 0.35 and any(skill.lower().replace(".", "") in normalized.lower().replace(".", "") for skill in KNOWN_INLINE_SKILLS)


def _looks_like_stop_section(line: str) -> bool:
    normalized = _normalize_key(line)
    return any(
        token in normalized
        for token in ("ingenieria", "tecnologia", "curso", "platzi", "daniel pinto", "agudelo", "jesus pavon", "jesus blanco")
    )


def _find_related_resume_files(path: Path, profile: ResumeProfile) -> list[Path]:
    directory = path.parent
    tokens = [token for token in _normalize_key(profile.full_name).split() if len(token) >= 4]
    related: list[Path] = []
    for candidate in directory.iterdir():
        if candidate == path or candidate.suffix.lower() not in {".pdf", ".docx"}:
            continue
        candidate_name = _normalize_key(candidate.stem)
        if "hv" not in candidate_name and "hoja de vida" not in candidate_name:
            continue
        matches = sum(1 for token in tokens if token in candidate_name)
        if matches >= 2:
            related.append(candidate)
    return related


def _merge_profiles(primary: ResumeProfile, secondary: ResumeProfile) -> ResumeProfile:
    return ResumeProfile(
        full_name=_pick_better_name(primary.full_name, secondary.full_name),
        location=_pick_better_location(primary.location, secondary.location),
        phone=_pick_better_text(primary.phone, secondary.phone),
        email=_pick_better_email(primary.email, secondary.email),
        linkedin=_pick_better_url(primary.linkedin, secondary.linkedin),
        github=_pick_better_url(primary.github, secondary.github),
        portfolio=_pick_better_url(primary.portfolio, secondary.portfolio),
        professional_summary=_pick_better_text(primary.professional_summary, secondary.professional_summary),
        technical_skills=_dedupe_keep_order(primary.technical_skills + secondary.technical_skills),
        soft_skills=_dedupe_keep_order(primary.soft_skills + secondary.soft_skills),
        experiences=_merge_experiences(primary.experiences, secondary.experiences),
        projects=_merge_projects(primary.projects, secondary.projects),
        education=_dedupe_keep_order(primary.education + secondary.education),
        certifications=_dedupe_keep_order(primary.certifications + secondary.certifications),
        languages=_dedupe_keep_order(primary.languages + secondary.languages),
        raw_text="\n".join(part for part in [primary.raw_text, secondary.raw_text] if part),
    )


def _sanitize_resume_profile(profile: ResumeProfile) -> ResumeProfile:
    education: list[str] = []
    certifications: list[str] = []
    for item in profile.education + profile.certifications:
        cleaned = _sanitize_profile_line(item)
        if not cleaned:
            continue
        normalized = _normalize_key(cleaned)
        if any(hint in normalized for hint in CERTIFICATION_HINTS):
            certifications.extend(_split_certification_line(cleaned))
        else:
            education.append(cleaned)

    languages = [
        _sanitize_profile_line(item)
        for item in profile.languages
        if _looks_like_language_line(item)
    ]
    languages = [item for item in languages if item]

    experiences = [_sanitize_experience(item) for item in profile.experiences]
    experiences = [item for item in experiences if item is not None]
    experiences = _normalize_and_merge_experiences(experiences)

    projects = [_sanitize_project(item) for item in profile.projects]
    projects = [item for item in projects if item is not None]

    summary = _sanitize_summary(profile.professional_summary)

    return ResumeProfile(
        full_name=_sanitize_profile_line(profile.full_name),
        location=_sanitize_profile_line(profile.location),
        phone=_normalize_phone(profile.phone),
        email=_normalize_email(profile.email),
        linkedin=_sanitize_url(profile.linkedin),
        github=_sanitize_url(profile.github),
        portfolio=_sanitize_url(profile.portfolio),
        professional_summary=summary,
        technical_skills=[item for item in (_sanitize_profile_line(skill) for skill in profile.technical_skills) if item],
        soft_skills=[item for item in (_sanitize_profile_line(skill) for skill in profile.soft_skills) if item],
        experiences=experiences,
        projects=projects,
        education=_dedupe_keep_order(education),
        certifications=_dedupe_keep_order(certifications),
        languages=_dedupe_keep_order(languages),
        raw_text=profile.raw_text,
    )


def _pick_better_text(primary: str, secondary: str) -> str:
    if not primary:
        return secondary
    if not secondary:
        return primary
    return primary if len(primary) >= len(secondary) else secondary


def _pick_better_name(primary: str, secondary: str) -> str:
    cleaned_primary = _strip_contact_noise_from_name(primary)
    cleaned_secondary = _strip_contact_noise_from_name(secondary)
    if cleaned_primary and len(cleaned_primary.split()) >= 3:
        return cleaned_primary
    if cleaned_secondary and len(cleaned_secondary.split()) >= 3:
        return cleaned_secondary
    return cleaned_primary or cleaned_secondary


def _pick_better_location(primary: str, secondary: str) -> str:
    cleaned_primary = _clean_location(primary)
    cleaned_secondary = _clean_location(secondary)
    if _looks_like_location(cleaned_primary):
        return cleaned_primary
    if _looks_like_location(cleaned_secondary):
        return cleaned_secondary
    return cleaned_primary or cleaned_secondary


def _pick_better_email(primary: str, secondary: str) -> str:
    cleaned_primary = _normalize_email(primary)
    cleaned_secondary = _normalize_email(secondary)
    if "@" in cleaned_primary:
        return cleaned_primary
    if "@" in cleaned_secondary:
        return cleaned_secondary
    return cleaned_primary or cleaned_secondary


def _pick_better_url(primary: str, secondary: str) -> str:
    cleaned_primary = _normalize_url(primary) if primary else ""
    cleaned_secondary = _normalize_url(secondary) if secondary else ""
    if cleaned_primary.startswith("http"):
        return cleaned_primary
    if cleaned_secondary.startswith("http"):
        return cleaned_secondary
    return cleaned_primary or cleaned_secondary


def _merge_experiences(primary: list[ResumeExperience], secondary: list[ResumeExperience]) -> list[ResumeExperience]:
    merged: list[ResumeExperience] = []
    index_by_key: dict[str, int] = {}
    for experience in [*primary, *secondary]:
        key = "|".join(
            [
                _normalize_key(experience.role),
                _normalize_key(experience.company),
                _normalize_key(experience.start_date),
                _normalize_key(experience.end_date),
            ]
        )
        if key in index_by_key:
            current = merged[index_by_key[key]]
            merged[index_by_key[key]] = ResumeExperience(
                role=_pick_better_text(current.role, experience.role),
                company=_pick_better_text(current.company, experience.company),
                location=_pick_better_text(current.location, experience.location),
                start_date=_pick_better_text(current.start_date, experience.start_date),
                end_date=_pick_better_text(current.end_date, experience.end_date),
                bullets=_dedupe_keep_order(current.bullets + experience.bullets),
            )
            continue
        merged.append(experience)
        index_by_key[key] = len(merged) - 1
    return merged


def _merge_projects(primary: list[ResumeProject], secondary: list[ResumeProject]) -> list[ResumeProject]:
    merged: list[ResumeProject] = []
    seen: set[str] = set()
    for project in [*primary, *secondary]:
        key = "|".join([_normalize_key(project.name), _normalize_key(project.role), _normalize_key(project.url)])
        if key in seen:
            continue
        merged.append(project)
        seen.add(key)
    return merged


def _strip_contact_noise_from_name(value: str) -> str:
    cleaned = re.split(r"linkedin\.com|github\.com|https?://|@|\d|[|·]", value, maxsplit=1, flags=re.IGNORECASE)[0]
    cleaned = cleaned.strip(" ,-")
    tokens = cleaned.split()
    if len(tokens) > 5:
        tokens = tokens[:5]
    return " ".join(tokens).strip()


def _sanitize_profile_line(value: str) -> str:
    cleaned = " ".join(value.strip().split())
    if not cleaned:
        return ""
    normalized = _normalize_key(cleaned)
    if any(fragment in normalized for fragment in BLOCKED_PROFILE_FRAGMENTS):
        return ""
    if normalized in {"formacion complementaria", "formación complementaria", "habilidades", "idiomas"}:
        return ""
    if len(cleaned) > 220:
        return ""
    return cleaned


def _sanitize_summary(value: str) -> str:
    cleaned = _sanitize_profile_line(value)
    if not cleaned:
        return ""
    normalized = _normalize_key(cleaned)
    if any(token in normalized for token in ("referencias", "mas informacion", "disponibilidad", "experiencia laboral", "formacion academica")):
        return ""
    if len(cleaned.split()) > 80:
        sentences = re.split(r"(?<=[.!?])\s+", cleaned)
        cleaned = " ".join(sentences[:2]).strip()
    return cleaned


def _looks_like_language_line(value: str) -> bool:
    normalized = _normalize_key(value)
    return any(hint in normalized for hint in LANGUAGE_HINTS)


def _split_certification_line(value: str) -> list[str]:
    parts = re.split(r"[•|]", value)
    results: list[str] = []
    for part in parts:
        cleaned = _sanitize_profile_line(part)
        if not cleaned:
            continue
        if cleaned.lower().startswith("formación complementaria") or cleaned.lower().startswith("formacion complementaria"):
            continue
        results.append(cleaned)
    return results


def _sanitize_experience(experience: ResumeExperience) -> ResumeExperience | None:
    role = _sanitize_profile_line(experience.role)
    company = _sanitize_profile_line(experience.company)
    if not role and not company:
        return None
    normalized_role = _normalize_key(role)
    normalized_company = _normalize_key(company)
    if "salitre" in normalized_role or "salitre" in normalized_company or normalized_role == "operador":
        return None
    bullets = []
    for bullet in experience.bullets:
        cleaned = _sanitize_bullet_line(bullet)
        if cleaned:
            bullets.append(cleaned)
    return ResumeExperience(
        role=role,
        company=company,
        location=_sanitize_profile_line(experience.location),
        start_date=_sanitize_profile_line(experience.start_date),
        end_date=_sanitize_profile_line(experience.end_date),
        bullets=_dedupe_keep_order(bullets),
    )


def _normalize_and_merge_experiences(experiences: list[ResumeExperience]) -> list[ResumeExperience]:
    merged: list[ResumeExperience] = []
    index_by_key: dict[str, int] = {}
    for experience in experiences:
        canonical = _canonicalize_experience(experience)
        key = "|".join([_normalize_key(canonical.company), _normalize_key(canonical.role)])
        if key in index_by_key:
            current = merged[index_by_key[key]]
            merged[index_by_key[key]] = ResumeExperience(
                role=_pick_better_text(current.role, canonical.role),
                company=_pick_better_text(current.company, canonical.company),
                location=_pick_better_text(current.location, canonical.location),
                start_date=_pick_earliest_date(current.start_date, canonical.start_date),
                end_date=_pick_latest_date(current.end_date, canonical.end_date),
                bullets=_dedupe_keep_order(current.bullets + canonical.bullets),
            )
            continue
        merged.append(canonical)
        index_by_key[key] = len(merged) - 1
    return merged


def _canonicalize_experience(experience: ResumeExperience) -> ResumeExperience:
    role = experience.role
    company = experience.company
    location = experience.location
    start_date = experience.start_date
    end_date = experience.end_date
    bullets = [_trim_sentence(bullet) for bullet in experience.bullets if _trim_sentence(bullet)]
    normalized_company = _normalize_key(company)
    normalized_role = _normalize_key(role)
    normalized_text = _normalize_key(" ".join([company, role, location, start_date, end_date, *bullets]))

    if "kepri" in normalized_company or "bootcamp udec" in normalized_role:
        role = "Desarrollador Frontend Web"
        company = "Kepri Holística"
        location = "Soacha, Colombia"
        start_date = "Octubre 2024"
        end_date = "Noviembre 2024"
        bullets = [
            bullet
            for bullet in bullets
            if any(
                token in _normalize_key(bullet)
                for token in ("bootcamp", "next.js", "e-commerce", "catalogo", "catálogo", "pagos", "responsivo", "plataforma")
            )
        ]
    elif "charles barber" in normalized_company:
        role = "Desarrollador Android"
        company = "Charles Barber"
        location = "Bogotá, Colombia"
        start_date = start_date or "Agosto 2024"
        end_date = end_date or "Octubre 2024"
    elif "chocontano restaurante" in normalized_company:
        role = "Desarrollador Frontend y Mobile"
        company = "Chocontano Restaurante"
        location = "Soacha, Colombia"
        start_date = start_date or "Marzo 2024"
        end_date = end_date or "Junio 2024"
    elif "universidad de cundinamarca fusagasuga" in normalized_company or "emprex360" in normalized_text:
        role = "Desarrollador Full Stack Web"
        company = "Emprex360"
        location = "Remoto"
        start_date = "Enero 2024"
        end_date = "Febrero 2025"
    elif normalized_company == "universidad de cundinamarca" and normalized_role == "auxiliar de oficina ii":
        location = "Soacha, Colombia"
    elif normalized_company == "universidad de cundinamarca" and "tecnico en soporte de equipos de computo" in normalized_role:
        location = "Soacha, Colombia"

    bullets = _clean_canonical_bullets(company, role, bullets)
    return ResumeExperience(
        role=role,
        company=company,
        location=location,
        start_date=start_date,
        end_date=end_date,
        bullets=_dedupe_keep_order(bullets),
    )


def _clean_canonical_bullets(company: str, role: str, bullets: list[str]) -> list[str]:
    normalized_company = _normalize_key(company)
    normalized_role = _normalize_key(role)
    filtered: list[str] = []
    for bullet in bullets:
        normalized = _normalize_key(bullet)
        if any(token in normalized for token in ("skills adicionales", "educacion y certificacion", "formacion complementario", "formacion complementaria")):
            continue
        if "kepri" in normalized_company and any(
            token in normalized
            for token in ("lactato", "deport", "ciencias del deporte", "datos biometricos", "datos biométricos")
        ):
            continue
        if normalized_role == "auxiliar de oficina ii" and not any(
            token in normalized
            for token in (
                "centros de computo",
                "centros de cómputo",
                "recursos tecnologicos",
                "recursos tecnológicos",
                "automatizacion",
                "automatización",
                "inventario",
                "reportes",
                "usuarios",
                "equipos",
            )
        ):
            continue
        filtered.append(bullet)
    return filtered


def _pick_earliest_date(primary: str, secondary: str) -> str:
    return _pick_boundary_date(primary, secondary, earliest=True)


def _pick_latest_date(primary: str, secondary: str) -> str:
    return _pick_boundary_date(primary, secondary, earliest=False)


def _pick_boundary_date(primary: str, secondary: str, earliest: bool) -> str:
    primary_parts = _parse_month_year(primary)
    secondary_parts = _parse_month_year(secondary)
    if primary_parts is None:
        return secondary
    if secondary_parts is None:
        return primary
    if earliest:
        return primary if primary_parts <= secondary_parts else secondary
    return primary if primary_parts >= secondary_parts else secondary


def _parse_month_year(value: str) -> tuple[int, int] | None:
    normalized = _normalize_key(value)
    match = re.search(r"(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+(\d{4})", normalized)
    if not match:
        return None
    return int(match.group(2)), MONTH_ORDER[match.group(1)]


def _sanitize_project(project: ResumeProject) -> ResumeProject | None:
    name = _sanitize_profile_line(project.name)
    role = _sanitize_profile_line(project.role)
    normalized = _normalize_key(" ".join([name, role]))
    if "salitre" in normalized:
        return None
    bullets = []
    for bullet in project.bullets:
        cleaned = _sanitize_bullet_line(bullet)
        if cleaned:
            bullets.append(cleaned)
    return ResumeProject(
        name=name,
        role=role,
        technologies=[item for item in (_sanitize_profile_line(tech) for tech in project.technologies) if item],
        bullets=_dedupe_keep_order(bullets),
        url=_sanitize_url(project.url),
    )


def _sanitize_bullet_line(value: str) -> str:
    cleaned = " ".join(value.strip().split())
    cleaned = BULLET_PREFIX_RE.sub("", cleaned).strip()
    if not cleaned:
        return ""
    cleaned = re.split(r"\bOperador\b", cleaned, maxsplit=1, flags=re.IGNORECASE)[0].strip()
    if not cleaned:
        return ""
    normalized = _normalize_key(cleaned)
    if any(fragment in normalized for fragment in BLOCKED_PROFILE_FRAGMENTS):
        return ""
    if any(token in normalized for token in ("referencias", "mas informacion", "disponibilidad", "formacion academica")):
        return ""
    if " operad" in normalized and "salitre" in normalized:
        return ""
    if "vehicul" in normalized or "parque" in normalized:
        return ""
    return cleaned


def _sanitize_url(value: str) -> str:
    cleaned = value.strip()
    if not cleaned:
        return ""
    normalized = _normalize_key(cleaned)
    if any(fragment in normalized for fragment in BLOCKED_PROFILE_FRAGMENTS):
        return ""
    return _normalize_url(cleaned)


def _normalize_phone(value: str) -> str:
    return " ".join(value.strip().split())


def _clean_location(value: str) -> str:
    cleaned = re.split(r"linkedin\.com|github\.com|https?://|@|\d|[|·]", value, maxsplit=1, flags=re.IGNORECASE)[0]
    cleaned = cleaned.strip(" ,-")
    return re.sub(r"\s+", " ", cleaned)


def _looks_like_location(value: str) -> bool:
    if not value:
        return False
    if len(value) > 60:
        return False
    if "@" in value or "http" in value.lower():
        return False
    return len(value.split()) <= 4


def _ends_with_connector(value: str) -> bool:
    stripped = value.strip().rstrip(".,;:")
    if not stripped:
        return False
    return stripped.split()[-1].lower() in CONNECTOR_ENDINGS


def _strip_trailing_connectors(value: str) -> str:
    words = value.strip().split()
    while words and words[-1].rstrip(".,;:").lower() in CONNECTOR_ENDINGS:
        words.pop()
    return " ".join(words).strip()


def _dedupe_keep_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        key = item.lower()
        if item and key not in seen:
            result.append(item)
            seen.add(key)
    return result
