from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT = Path("JobOps_Architecture_and_Documentation.docx")


def configure_styles(document: Document) -> None:
    styles = document.styles
    for style_name, size in (
        ("Normal", 10.5),
        ("Title", 20),
        ("Subtitle", 11),
        ("Heading 1", 15),
        ("Heading 2", 12.5),
        ("Heading 3", 11),
    ):
        try:
            style = styles[style_name]
        except KeyError:
            continue
        style.font.name = "Calibri"
        style.font.size = Pt(size)


def add_paragraph(document: Document, text: str = "", *, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_mono_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    for index, line in enumerate(text.strip("\n").splitlines()):
        if index:
            run.add_break()
        run.add_text(line)


def add_table(document: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_document() -> Document:
    document = Document()
    configure_styles(document)

    document.core_properties.title = "JobOps Architecture and Documentation"
    document.core_properties.subject = "JobOps Personal Assistant technical documentation"
    document.core_properties.author = "OpenAI Codex"
    document.core_properties.comments = "Basado exclusivamente en el codigo real del repositorio."

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("JobOps Assistant\nArchitecture and Technical Documentation")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Documento tecnico basado exclusivamente en el codigo real del repositorio actual. "
        f"Generado el {datetime.now().strftime('%Y-%m-%d %H:%M')}.\n"
        "Alcance: CLI, scrapers, monitor, fuentes, persistencia SQLite, filtrado, Telegram y generacion de CV."
    )

    document.add_heading("1. Introduccion del Proyecto", level=1)
    add_paragraph(
        document,
        "JobOps Personal Assistant es una aplicacion local en Python, orientada a linea de comandos, "
        "para organizar una busqueda laboral tecnica. El sistema permite registrar un perfil del candidato, "
        "definir fuentes de busqueda por portal, ejecutar scrapers sobre URLs publicas, filtrar ofertas por "
        "relevancia, guardar resultados en SQLite y emitir alertas por Telegram.",
    )
    add_paragraph(
        document,
        "El problema que resuelve es operativo: centralizar resultados dispersos de multiples portales, reducir "
        "ruido mediante reglas de matching y conservar evidencia persistente de lo encontrado, descartado, "
        "notificado o duplicado. El proyecto combina scraping responsable, evaluacion heuristica y seguimiento "
        "manual de postulaciones, sin autoaplicar y sin resolver captchas.",
    )
    add_bullets(
        document,
        [
            "Entrada principal: comandos de `main.py` que delegan a `src/jobops_assistant/cli.py`.",
            "Persistencia: SQLite via SQLAlchemy en `database.py` y `models.py`.",
            "Scraping: base requests+BeautifulSoup y variantes Selenium para portales dinamicos o bloqueados.",
            "Filtrado: `matcher.py` y `discarded_job_service.py` deciden relevancia, score y razones de descarte.",
            "Alertas: `telegram_notifier.py` envia mensajes individuales o digest y registra la entrega en BD.",
        ],
    )
    add_paragraph(
        document,
        "Importante: la filosofia observable del codigo es conservadora. El README y las excepciones del "
        "scraping dejan claro que el sistema evita login automation, bypass de captcha, proxies y autoaplicacion.",
    )

    document.add_heading("2. Arquitectura General", level=1)
    document.add_heading("2.1 Componentes Principales", level=2)
    add_bullets(
        document,
        [
            "CLI: `main.py` solo redirige a `cli.py`; `cli.py` define parser, subcomandos y handlers.",
            "Settings: `settings.py` carga `.env` y materializa la configuracion en el dataclass `Settings`.",
            "Sources Manager: `search_sources.py` crea, prueba, habilita, pausa, reanuda y agenda fuentes.",
            "Scraper Registry: `scrapers/registry.py` mapea cada portal a su clase concreta.",
            "Scrapers: `scrapers/base_scraper.py` y `scrapers/selenium_base.py` proveen las bases; los portales concretos viven en `scrapers/*.py`.",
            "Monitor Engine: `freshness_monitor.py` ejecuta el ciclo de lectura, descarte, deduplicacion, guardado y notificacion.",
            "Filter Engine: `matcher.py` calcula score y `discarded_job_service.py` decide si la oferta se guarda o se audita como descartada.",
            "Database Layer: `models.py`, `database.py`, `job_service.py`, `profile_service.py` y `resume_profile_service.py` administran entidades y operaciones.",
            "Notification System: `telegram_notifier.py` formatea mensajes, manda digest y registra resultados en `notifications`.",
            "Resume/CV Pipeline: `resume_reader.py`, `resume_profile_service.py`, `ats_resume_builder.py` y `cv_generator.py` cubren importacion de HV y generacion de CVs.",
            "Daily Workflow: `workflows.py` integra Gmail, parser de alertas y Telegram; `gmail_reader.py` hoy es un stub conservador.",
        ],
    )

    document.add_heading("2.2 Diagrama Textual de Arquitectura", level=2)
    add_mono_block(
        document,
        """Usuario
  |
  v
main.py
  |
  v
cli.py -------------------------------------------------------------+
  |                                                                 |
  +--> settings.py                                                  |
  +--> database.py / models.py                                      |
  +--> search_sources.py <--> tabla job_search_sources              |
  +--> profile_service.py <--> tabla candidate_profile              |
  +--> freshness_monitor.py -------------------------------------+  |
  |                                                             |  |
  |                                                             v  |
  |                                                   scrapers/registry.py
  |                                                             |
  |                           +-----------------+---------------+----------------+
  |                           |                 |                                |
  |                           v                 v                                v
  |                  LinkedIn Selenium   Indeed Selenium                Computrabajo Selenium
  |                           |                 |                                |
  |                           +---------> lista de ScrapedJob <------------------+
  |                                               |
  |                                               +--> discarded_job_service.py --> tabla discarded_jobs
  |                                               |
  |                                               +--> job_service.py + matcher.py
  |                                                          |
  |                                                          +--> tabla job_offers
  |                                                          +--> tabla job_seen_hashes
  |                                                          +--> tabla generated_documents
  |                                                          +--> tabla notifications
  |                                                          |
  |                                                          v
  +--------------------------------------------------> telegram_notifier.py --> Telegram Bot API""",
    )

    document.add_heading("2.3 Flujo End-to-End", level=2)
    add_paragraph(document, "Flujo solicitado por el usuario y confirmado por el codigo:")
    add_mono_block(document, "User -> CLI -> Sources -> Scraper -> Filter -> DB -> Telegram")
    add_paragraph(
        document,
        "Ese flujo se implementa asi: el usuario registra o prueba una fuente en `cli.py`; "
        "`search_sources.py` persiste la configuracion; `freshness_monitor.py` solicita al registry el "
        "scraper correcto; el scraper devuelve objetos `ScrapedJob`; `discarded_job_service.py` y `matcher.py` "
        "deciden si la oferta se conserva; `job_service.py` la guarda o fusiona con una existente; finalmente "
        "`telegram_notifier.py` envia digest si la oferta supera el umbral y aun no fue notificada.",
    )

    document.add_heading("2.4 Scheduler e Intervalos", level=2)
    add_paragraph(
        document,
        "No existe un scheduler externo tipo cron embebido ni un worker distribuido. El agendamiento es local y simple: "
        "`monitor watch` ejecuta un bucle `while True`, abre una sesion SQLAlchemy por ciclo, llama a "
        "`run_fresh_monitor(..., force_all=False, notify_pending=True)` y luego duerme el numero de minutos indicado. "
        "La decision de si una fuente esta due se calcula en `search_sources.get_due_sources()` usando "
        "`last_checked_at + interval_minutes`.",
    )

    document.add_heading("3. Estructura de Carpetas y Modulos Reales", level=1)
    add_paragraph(
        document,
        "Nota importante de arquitectura: el codigo real no esta organizado fisicamente en subdirectorios "
        "`core/`, `monitor/`, `database/`, `notifications/` o `utils/`. Esas responsabilidades existen de forma "
        "logica, pero el paquete real es mayormente plano dentro de `src/jobops_assistant/`, con un unico "
        "subpackage dedicado a scrapers.",
    )

    document.add_heading("3.1 Estructura del Repositorio", level=2)
    add_bullets(
        document,
        [
            "`main.py`: entry point minimo que delega al package principal.",
            "`src/`: codigo fuente principal.",
            "`tests/`: suite de pruebas unitarias y funcionales de bajo nivel.",
            "`data/`: artefactos locales; en esta copia contiene BD SQLite, `resume_profile.json` y perfiles de navegador usados en pruebas Selenium.",
            "`generated/`: salidas generadas por el sistema, especialmente CVs en `generated/cvs/`.",
            "`debug/`: snapshots HTML y metadatos cuando se usa `sources test --debug-html`.",
            "`templates/`: carpeta reservada para plantillas; actualmente solo contiene `.gitkeep`.",
            "`config/`: carpeta placeholder; actualmente solo contiene `.gitkeep`.",
            "`.env.example`: valores de ejemplo para la configuracion.",
            "`update_linkedin_interval.py`: script auxiliar puntual para mantenimiento manual de intervalos en SQLite.",
        ],
    )

    document.add_heading("3.2 Arbol Real del Package Principal", level=2)
    add_mono_block(
        document,
        """src/jobops_assistant/
|-- application_types.py
|-- ats_resume_builder.py
|-- cli.py
|-- cv_generator.py
|-- database.py
|-- date_utils.py
|-- discarded_job_service.py
|-- freshness_monitor.py
|-- gmail_reader.py
|-- job_parser.py
|-- job_service.py
|-- linkedin_url.py
|-- matcher.py
|-- message_generator.py
|-- models.py
|-- profile_service.py
|-- resume_profile_service.py
|-- resume_reader.py
|-- schemas.py
|-- search_sources.py
|-- settings.py
|-- telegram_notifier.py
|-- workflows.py
`-- scrapers/
    |-- base_scraper.py
    |-- computrabajo_scraper.py
    |-- elempleo_scraper.py
    |-- getonboard_scraper.py
    |-- indeed_scraper.py
    |-- indeed_selenium_scraper.py
    |-- linkedin_scraper.py
    |-- linkedin_selenium_scraper.py
    |-- magneto_scraper.py
    |-- registry.py
    |-- selenium_base.py
    |-- sena_scraper.py
    `-- torre_scraper.py""",
    )

    document.add_heading("3.3 Inventario de Modulos del Package", level=2)
    add_table(
        document,
        ["Archivo", "Dominio", "Responsabilidad real"],
        [
            ("application_types.py", "Constantes compartidas", "Declara los tipos `linkedin_easy_apply`, `external_apply` y `unknown`."),
            ("ats_resume_builder.py", "CV ATS", "Construye CVs ATS por target a partir de un `ResumeProfile` y opcionalmente una `JobOffer`."),
            ("cli.py", "CLI", "Define parser, subcomandos y handlers de toda la operacion del sistema."),
            ("cv_generator.py", "CV simple", "Genera un DOCX simple a partir de `CandidateProfile` y una oferta concreta."),
            ("database.py", "Bootstrap de BD", "Crea engine/session factory, inicializa tablas y ejecuta migraciones SQLite ligeras."),
            ("date_utils.py", "Fechas", "Normaliza timezones, interpreta textos relativos y da formato para mensajes."),
            ("discarded_job_service.py", "Descartadas", "Analiza relevancia, guarda descartadas, exporta y reprocesa."),
            ("freshness_monitor.py", "Monitor", "Ejecuta scraping por ciclo, deduplica, guarda, arma digest y registra notificaciones."),
            ("gmail_reader.py", "Gmail", "Adaptador preparado para leer alertas; hoy devuelve vacio de forma conservadora."),
            ("job_parser.py", "Parser textual", "Extrae un `ParsedJobOffer` desde texto libre con patrones simples."),
            ("job_service.py", "CRUD de ofertas", "Crea, fusiona, lista, limpia ofertas y controla marca de Telegram."),
            ("linkedin_url.py", "Normalizacion LinkedIn", "Limpia URLs de busqueda y fuerza parametros canonicos de LinkedIn."),
            ("matcher.py", "Matching y relevancia", "Calcula score y aplica reglas de relevancia por target."),
            ("message_generator.py", "Mensajes", "Crea mensajes breves de postulacion usando perfil y oferta."),
            ("models.py", "ORM schema", "Declara todas las entidades SQLAlchemy del proyecto."),
            ("profile_service.py", "Perfil candidato", "Obtiene o actualiza el perfil unico del candidato."),
            ("resume_profile_service.py", "Persistencia de HV", "Guarda y carga `ResumeProfile` en JSON local."),
            ("resume_reader.py", "Importacion de HV", "Lee DOCX/PDF, limpia texto y construye `ResumeProfile`."),
            ("schemas.py", "Dataclasses", "Modelos auxiliares para matching, parsing y resume data."),
            ("search_sources.py", "Fuentes", "Gestiona fuentes, intervalos, pausa, prueba y due scheduling."),
            ("settings.py", "Configuracion", "Carga variables de entorno y las agrupa en `Settings`."),
            ("telegram_notifier.py", "Notificaciones", "Formatea y envia alertas/digest a Telegram y registra entregas."),
            ("workflows.py", "Workflow diario", "Orquesta la lectura de alertas Gmail y notificacion diaria."),
        ],
    )

    document.add_heading("3.4 Inventario del Subpackage `scrapers/`", level=2)
    add_table(
        document,
        ["Archivo", "Tipo", "Responsabilidad real"],
        [
            ("base_scraper.py", "Base requests", "Pipeline comun con requests, BeautifulSoup, normalizacion de URLs y excepciones de bloqueo/login/captcha."),
            ("selenium_base.py", "Base Selenium", "Obtiene HTML renderizado con Chrome y aplica scroll, opciones de perfil y deteccion base de bloqueo."),
            ("registry.py", "Registry", "Relaciona el nombre del portal con su clase scraper."),
            ("linkedin_scraper.py", "LinkedIn estatico", "Parser sencillo para HTML publico de LinkedIn sin navegador real."),
            ("linkedin_selenium_scraper.py", "LinkedIn Selenium", "Scraper principal de LinkedIn con cards publicas/logueadas, deteccion de authwall y filtro de tipo de solicitud."),
            ("indeed_scraper.py", "Indeed estatico", "Scraper por selectores sobre HTML publico simple."),
            ("indeed_selenium_scraper.py", "Indeed Selenium", "Scraper browser-based que detecta `jk` y canoniza URLs `viewjob`."),
            ("computrabajo_scraper.py", "Computrabajo Selenium", "Scraper browser-based con extraccion robusta multimetodo y estados `blocked/no_results/scraper_broken/ok`."),
            ("elempleo_scraper.py", "Elempleo", "Scraper por selectores con heuristicas de titulo, empresa, ubicacion y fecha."),
            ("magneto_scraper.py", "Magneto", "Scraper selector-based simple."),
            ("torre_scraper.py", "Torre", "Scraper selector-based simple."),
            ("getonboard_scraper.py", "GetOnBoard", "Scraper selector-based simple."),
            ("sena_scraper.py", "SENA", "Scraper selector-based simple."),
        ],
    )

    document.add_heading("4. Scrapers", level=1)
    document.add_heading("4.1 Base Comun de Scraping", level=2)
    add_paragraph(
        document,
        "`scrapers/base_scraper.py` define el contrato del subsistema. El dataclass `ScrapedJob` representa una "
        "oferta ya parseada con titulo, empresa, portal, ubicacion, modalidad, salario, URL, descripcion, "
        "requirements, fechas y `application_type`. `ResponseDebugSnapshot` conserva evidencia de URL pedida, "
        "status code, URL final, content type, HTML y razon de bloqueo para depuracion.",
    )
    add_bullets(
        document,
        [
            "`BaseJobScraper.fetch_search_results()` hace la lectura base y llena `last_response_debug`.",
            "`BaseJobScraper.scrape()` ejecuta el ciclo comun: obtener HTML -> parsear cards -> normalizar URL -> leer detalle -> limitar por `max_results_per_source`.",
            "`normalize_url()` elimina parametros de tracking (`trk`, `utm_*`, `currentjobid`, etc.).",
            "`_request_text()` trata `403` como bloqueo, `429` como rate limit y luego aplica deteccion por contenido.",
            "`SelectorBasedScraper` reutiliza selectores CSS para titulo, empresa, ubicacion, fecha, descripcion y requirements.",
        ],
    )

    document.add_heading("4.2 Base Selenium", level=2)
    add_paragraph(
        document,
        "`scrapers/selenium_base.py` extiende el modelo anterior para obtener HTML renderizado con Selenium y Chrome. "
        "La clase `SeleniumJobScraper` usa `webdriver.Chrome`, aplica opciones de navegador, espera una carga basica, "
        "hace scroll y finalmente parsea `driver.page_source` con BeautifulSoup.",
    )
    add_bullets(
        document,
        [
            "Respeta `JOBOPS_ENABLE_SELENIUM`; si esta en `false`, los scrapers Selenium no operan.",
            "Soporta `JOBOPS_SELENIUM_HEADLESS`, `JOBOPS_SELENIUM_PAGE_LOAD_TIMEOUT`, `JOBOPS_SELENIUM_SCROLL_PAUSE` y `JOBOPS_SELENIUM_MAX_SCROLLS`.",
            "Soporta reutilizar perfil local de Chrome con `JOBOPS_SELENIUM_USER_DATA_DIR` y `JOBOPS_SELENIUM_PROFILE_DIRECTORY`.",
            "Inyecta `JOBOPS_SCRAPER_USER_AGENT` en Chrome.",
            "Aplica deteccion base de captcha, recaptcha, turnstile, access denied y login wall por texto y selectores.",
        ],
    )
    add_paragraph(
        document,
        "Salvo que una subclase administre el driver manualmente, la implementacion base abre y cierra una sesion "
        "Chrome por llamada de scraping. Computrabajo es la excepcion relevante: reutiliza el mismo driver durante "
        "listado y detalle dentro de `scrape()`.",
    )

    document.add_heading("4.3 Indeed Selenium", level=2)
    add_paragraph(
        document,
        "`indeed_selenium_scraper.py` es un scraper browser-based relativamente compacto. Su objetivo principal no es "
        "leer detalle profundo sino reconstruir una URL canonica estable por oferta. Para ello extrae el identificador "
        "`jk` desde `data-jk`, nodos internos con `data-jk` o desde los hrefs de Indeed.",
    )
    add_bullets(
        document,
        [
            "Si no logra extraer `jk`, la card se descarta.",
            "Cada oferta valida se normaliza a `https://co.indeed.com/viewjob?jk=...`.",
            "Deduplica por `jk` antes de producir resultados.",
            "La deteccion de captcha o Security Check se hereda del framework Selenium base.",
        ],
    )

    document.add_heading("4.4 LinkedIn Selenium", level=2)
    add_paragraph(
        document,
        "`linkedin_selenium_scraper.py` es el scraper mas sofisticado del proyecto. Puede trabajar sobre cards "
        "publicas y tambien sobre layouts visibles cuando el usuario reutiliza un perfil local de Chrome ya logueado. "
        "No automatiza el login: el soporte real es reusar sesion del navegador via Selenium settings.",
    )
    document.add_heading("4.4.1 Construccion de URL", level=3)
    add_paragraph(
        document,
        "La funcion `build_linkedin_jobs_url()` recibe keyword, location, date filter, experience levels y workplace "
        "types. Traduce esos filtros a parametros de LinkedIn (`f_TPR`, `f_E`, `f_WT`) y luego delega a "
        "`linkedin_url.build_linkedin_url()` para canonizar la URL final.",
    )
    add_bullets(
        document,
        [
            "`linkedin_url.build_linkedin_url()` elimina `currentJobId` y `origin`.",
            "Si encuentra `sortBy=R` u otro valor distinto de `DD`, lo reemplaza.",
            "Mantiene `f_TPR` si ya existe o si fue construido desde `date_posted`.",
            "Fuerza siempre `f_AL=true` y `sortBy=DD`.",
            "Registra por consola: URL original, URL limpia y parametros eliminados.",
        ],
    )
    document.add_heading("4.4.2 Extraccion y Navegacion", level=3)
    add_bullets(
        document,
        [
            "Usa multiples selectores para cards publicas y cards visibles en modo logueado.",
            "Hace scroll limitado y trata de pulsar `.infinite-scroller__show-more-button` cuando aparece.",
            "Intenta cerrar el modal de sign-in si LinkedIn muestra el overlay publico.",
            "Puede enriquecer descripcion desde el panel visible o, si `JOBOPS_LINKEDIN_FETCH_DETAILS=true`, abrir cada oferta y leer `div.description__text`.",
        ],
    )
    document.add_heading("4.4.3 Deteccion de Bloqueo y Vacio Real", level=3)
    add_bullets(
        document,
        [
            "Si hay cards publicas o cards visibles del layout logueado, el scraper considera la pagina valida.",
            "Si el HTML contiene mensajes de no results found o equivalentes, el scraper no pausa la fuente; simplemente devuelve cero resultados.",
            "Si detecta `authwall`, `checkpoint`, `captcha`, `security check`, `join linkedin` u otras senales de login/bloqueo sin contenido publico, lanza excepciones de bloqueo/login.",
            "Diferencia explicitamente entre 0 resultados reales y no hay cards porque LinkedIn cerro el acceso publico.",
        ],
    )
    document.add_heading("4.4.4 Filtrado de Ofertas", level=3)
    add_paragraph(
        document,
        "El scraper detecta `application_type` inspeccionando texto y controles de la card o del panel de detalle. "
        "Los valores posibles son `linkedin_easy_apply`, `external_apply` y `unknown`. Luego, fuera del scraper, "
        "`discarded_job_service.analyze_linkedin_application_type_for_discard()` puede descartar cualquier oferta "
        "que no sea Easy Apply si `JOBOPS_LINKEDIN_ONLY_EASY_APPLY=true`.",
    )

    document.add_heading("4.5 LinkedIn estatico", level=2)
    add_paragraph(
        document,
        "`linkedin_scraper.py` sigue existiendo como variante simple basada en requests y selectores CSS. "
        "Sirve para HTML publico basico, pero su tolerancia al bloqueo es mucho menor que la del scraper Selenium. "
        "Su deteccion es binaria: si no hay cards y aparecen tokens de captcha/login, lanza excepcion; en otros casos "
        "termina marcando la fuente como no accesible publicamente.",
    )

    document.add_heading("4.6 Computrabajo Selenium", level=2)
    add_paragraph(
        document,
        "El estado actual del repositorio muestra que `computrabajo_scraper.py` ya esta migrado a Selenium y no "
        "depende del pipeline requests clasico para la lectura real de resultados. La clase hereda de "
        "`SeleniumJobScraper`, pero reimplementa `scrape()`, `fetch_search_results()` y `fetch_job_detail()` para "
        "reutilizar un mismo driver durante el ciclo completo y para evaluar con mayor precision los falsos positivos "
        "de bloqueo o de 0 ofertas.",
    )
    document.add_heading("4.6.1 Flujo Browser-Based", level=3)
    add_bullets(
        document,
        [
            "Abre la URL con Selenium, espera DOM renderizado, hace scroll y usa `driver.page_source` como fuente final.",
            "Registra logs como `[computrabajo] selenium_open_url=...`, `[computrabajo] page_loaded`, `[computrabajo] html_length=...` y `[computrabajo] extracting_jobs_from_rendered_html`.",
            "Reutiliza el driver en `_active_driver` para que la lectura de detalle no abra un segundo navegador innecesario dentro del mismo `scrape()`.",
            "Si hay timeout, `_load_rendered_html()` reintenta una vez. Si hubo carga parcial con contenido util, puede aceptar el HTML parcial.",
        ],
    )
    document.add_heading("4.6.2 Extraccion Robusta de Cards", level=3)
    add_paragraph(
        document,
        "La funcion `extract_job_cards_robust()` evita depender de un unico selector fragil. Prueba cuatro caminos: "
        "selectores primarios, selectores alternativos, una busqueda XPath generica por enlaces de oferta y un "
        "fallback por heuristica de links que contienen patrones como `/oferta-de-trabajo`, `/trabajo-`, `/oferta/` o `/ofertas/`. "
        "Luego deduplica por href normalizado.",
    )
    add_bullets(
        document,
        [
            "Loguea cantidades por estrategia: `selector=primary`, `selector=alt`, `selector=xpath`, `selector=fallback`.",
            "Nunca depende solo de una clase CSS puntual del portal.",
            "La deduplicacion interna evita repetir cards por URL en la misma pagina.",
        ],
    )
    document.add_heading("4.6.3 Deteccion de Estado", level=3)
    add_paragraph(
        document,
        "`detect_blocking_state(response, html, job_cards)` devuelve uno de cuatro estados: `blocked`, "
        "`scraper_broken`, `no_results` u `ok`.",
    )
    add_bullets(
        document,
        [
            "`blocked`: status 403/429, texto de captcha/acceso denegado, login wall o redirect de bloqueo.",
            "`no_results`: no hay job links relevantes y el HTML muestra senales explicitas de pagina valida sin resultados.",
            "`scraper_broken`: no hay job links relevantes y tampoco hay evidencia clara de sin resultados; se interpreta como cambio de DOM o carga defectuosa, no como bloqueo.",
            "`ok`: existen cards o links de empleo relevantes.",
        ],
    )
    add_paragraph(
        document,
        "Solo el estado `blocked` dispara excepciones que hacen subir `failure_count` y eventualmente pausar la fuente. "
        "Los estados `scraper_broken` o `no_results` dejan la fuente activa.",
    )
    document.add_heading("4.6.4 Enriquecimiento de Detalle", level=3)
    add_paragraph(
        document,
        "Al abrir el detalle de una oferta, el scraper intenta mezclar JSON-LD tipo `JobPosting` y contenido HTML. "
        "Puede extraer titulo, empresa, ubicacion, salario, fecha publicada, descripcion, requirements y modalidad "
        "desde `script[type='application/ld+json']`, `.box_detail`, listas de requisitos y etiquetas de salario.",
    )

    document.add_heading("4.7 Otros Scrapers Soportados", level=2)
    add_paragraph(
        document,
        "Ademas de LinkedIn, Indeed y Computrabajo, el registry expone scrapers para `elempleo`, `magneto`, `torre`, "
        "`getonboard` y `sena`. En el codigo actual todos ellos son variantes selector-based sobre HTML publico, "
        "sin una capa Selenium dedicada propia.",
    )
    add_bullets(
        document,
        [
            "`elempleo_scraper.py` es el mas robusto de este grupo: tiene heuristicas adicionales para titulo, empresa, ubicacion y fecha.",
            "`magneto_scraper.py`, `torre_scraper.py`, `getonboard_scraper.py` y `sena_scraper.py` son implementaciones selector-based sencillas.",
            "Todos estos scrapers se benefician del pipeline comun de normalizacion de URLs, parseo de fechas y deteccion basica de bloqueo.",
        ],
    )

    document.add_heading("5. Sistema de Fuentes (Sources)", level=1)
    add_paragraph(
        document,
        "Una fuente representa una URL de busqueda persistente que el monitor puede revisar periodicamente. "
        "El modelo ORM real es `JobSearchSource` y la tabla fisica se llama `job_search_sources`; no existe una "
        "tabla llamada literalmente `sources`.",
    )
    document.add_heading("5.1 Que Guarda una Fuente", level=2)
    add_bullets(
        document,
        [
            "Portal (`portal`).",
            "Rol objetivo (`target_role`).",
            "URL de busqueda (`search_url`).",
            "Keywords y location auxiliares.",
            "Flag de habilitacion (`enabled`).",
            "Intervalo en minutos (`interval_minutes`).",
            "Ultima revision, ultimo error, contador de fallos, ultimo fallo y `paused_until`.",
        ],
    )
    document.add_heading("5.2 Ciclo de Vida", level=2)
    add_bullets(
        document,
        [
            "`add_source()` valida el intervalo minimo y normaliza URLs de LinkedIn antes de guardar.",
            "`list_sources()` y `get_source_by_id()` permiten inspeccion.",
            "`set_source_enabled()` activa o desactiva la fuente.",
            "`update_source_interval()` y `update_portal_source_intervals()` cambian frecuencia.",
            "`unpause_source_by_id()` y `unpause_sources_by_portal()` limpian el estado de pausa.",
            "`disable_blocked_sources()` deshabilita fuentes que acumularon demasiados fallos reales.",
        ],
    )
    document.add_heading("5.3 Scheduling y Pausas", level=2)
    add_paragraph(
        document,
        "`search_sources.py` define `AUTO_PAUSE_FAILURE_THRESHOLD = 3` y `AUTO_PAUSE_DURATION = 24 horas`. "
        "Cuando el monitor captura `CaptchaRequiredError`, `LoginRequiredError` o `SourceBlockedError`, llama a "
        "`record_source_failure()`. Al tercer fallo, la fuente queda pausada hasta `failed_at + 24h`.",
    )
    add_paragraph(
        document,
        "En cambio, los errores genericos del scraper o los casos de `scraper_broken` no pasan por ese camino: "
        "se registran en `last_error`, pero no se interpretan como bloqueo publico del portal.",
    )
    document.add_heading("5.4 Deduccion de Duplicados", level=2)
    add_paragraph(
        document,
        "La deduplicacion ocurre en varias capas. Cada oferta se normaliza por URL; luego se calcula `url_hash` "
        "con SHA-256; ademas, `freshness_monitor.find_possible_duplicate()` busca coincidencias por titulo, empresa "
        "y portal. Los hashes vistos se persisten en `job_seen_hashes`, lo que evita reprocesar indefinidamente la "
        "misma publicacion.",
    )

    document.add_heading("6. Motor de Filtrado y Relevancia", level=1)
    document.add_heading("6.1 Dos Capas Distintas", level=2)
    add_paragraph(
        document,
        "El proyecto separa relevancia y score. Primero decide si la oferta merece conservarse como oferta valida "
        "o si debe entrar en `discarded_jobs`. Despues, para las ofertas validas, calcula un `compatibility_score` "
        "de 0 a 100 y genera razones de coincidencia.",
    )
    document.add_heading("6.2 Analisis de Relevancia por Target", level=2)
    add_paragraph(
        document,
        "`matcher.analyze_relevance_for_target()` concentra las reglas de descarte. Usa conjuntos amplios de "
        "senales: seniority, entrada, roles administrativos, falsos positivos frontend, senales de soporte TI, "
        "y reglas por target como `backend_junior`, `frontend_junior`, `fullstack_junior`, `devops_trainee`, "
        "`soporte_aplicaciones`, `infraestructura_junior`, `cloud_support` y `qa_junior`.",
    )
    add_bullets(
        document,
        [
            "Descarta titles claramente senior o administrativos cuando no corresponden al target.",
            "Detecta señales tecnicas requeridas por rol.",
            "Tolera ciertos textos senior en descripcion secundaria si el title sigue siendo de entrada, pero penaliza el score.",
            "Devuelve `RelevanceAnalysis` con `relevant`, `reasons` y `detected_keywords`.",
        ],
    )
    document.add_heading("6.3 Score de Compatibilidad", level=2)
    add_paragraph(
        document,
        "`matcher.calculate_match()` arranca desde una base de 35 puntos y luego suma o resta segun keywords globales, "
        "negative rules y coincidencia con targets inferidos desde el perfil o desde el texto de la oferta.",
    )
    add_bullets(
        document,
        [
            "Keywords globales positivas: `junior`, `trainee`, `entry level`, `sin experiencia`, `remoto`, `hibrido`, etc.",
            "Reglas negativas: `senior`, `semi senior`, `lider`, `arquitecto`, `mas de 3 anos`, ingles avanzado, etc.",
            "Reglas por target con peso y listas de tecnologias/terminos.",
            "Bonus por skills compartidas entre el perfil del candidato y el texto de la oferta.",
        ],
    )
    add_paragraph(
        document,
        "El resultado es un `MatchResult(score, reasons)` que luego se serializa en `JobOffer.compatibility_score` "
        "y `JobOffer.match_reason`.",
    )
    document.add_heading("6.4 Razones de Descarte", level=2)
    add_paragraph(
        document,
        "Cuando una oferta no supera el filtro de relevancia, `discarded_job_service.analyze_scraped_job_for_discard()` "
        "construye un `DiscardedJobReview` con la oferta original, las razones de descarte, keywords detectadas y "
        "un score preliminar solo para auditoria. Esa informacion se guarda en `discarded_jobs` y puede reexportarse "
        "o reprocesarse mas adelante.",
    )
    document.add_heading("6.5 Filtro Especial de LinkedIn Easy Apply", level=2)
    add_paragraph(
        document,
        "Si `JOBOPS_LINKEDIN_ONLY_EASY_APPLY=true`, el pipeline agrega un descarte portal-especifico: cualquier "
        "oferta de `linkedin_selenium` cuyo `application_type` no sea `linkedin_easy_apply` se guarda como descartada "
        "con la razon `no es solicitud sencilla de LinkedIn`.",
    )

    document.add_heading("7. Base de Datos SQLite", level=1)
    add_paragraph(
        document,
        "La persistencia se implementa con SQLAlchemy sobre SQLite. `database.py` crea el engine, fabrica sesiones "
        "y ejecuta `Base.metadata.create_all(engine)`. Despues corre `_apply_sqlite_migrations()` para agregar columnas "
        "faltantes en instalaciones existentes y para normalizar URLs de fuentes LinkedIn ya guardadas.",
    )
    document.add_heading("7.1 Tablas Reales", level=2)
    add_table(
        document,
        ["Tabla", "Rol", "Campos o notas clave"],
        [
            ("candidate_profile", "Perfil unico del candidato", "Nombre, email, telefono, ciudad, summary, skills, projects, education, target_roles."),
            ("job_offers", "Ofertas aceptadas", "Titulo, empresa, portal, location, modality, salary, url, description, requirements, score, reason, fechas, normalized_url, url_hash, source_id, flags Telegram."),
            ("job_search_sources", "Fuentes configuradas", "Portal, target_role, search_url, keywords, location, enabled, interval, failure_count, paused_until, last_error."),
            ("discarded_jobs", "Auditoria de descartes", "Oferta descartada, target_role, razones, keywords detectadas, score preliminar, source_url, seen_count, last_seen_at."),
            ("job_seen_hashes", "Memoria de vistos", "url_hash, normalized_url, portal, first_seen_at, last_seen_at."),
            ("generated_documents", "Documentos generados", "job_offer_id, doc_type, file_path, created_at."),
            ("notifications", "Historial de alertas", "job_offer_id, channel, status, message, created_at."),
        ],
    )
    document.add_heading("7.2 Relaciones Practicas", level=2)
    add_bullets(
        document,
        [
            "`JobOffer.source_id` enlaza la oferta con la fuente que la produjo.",
            "`DiscardedJob.source_id` enlaza el descarte con la fuente original.",
            "`Notification.job_offer_id` registra intentos de envio por canal.",
            "`GeneratedDocument.job_offer_id` vincula CVs generados con una oferta.",
            "`JobOffer.url` tiene unique constraint y `DiscardedJob.url_hash` tambien es unico.",
        ],
    )
    document.add_heading("7.3 Nota sobre alerts", level=2)
    add_paragraph(
        document,
        "En la peticion se menciona una tabla alerts. En el codigo real no existe una tabla con ese nombre. "
        "La responsabilidad equivalente esta materializada en la tabla `notifications` y en los campos "
        "`telegram_notified` / `telegram_notified_at` de `job_offers`.",
    )

    document.add_heading("8. Telegram Bot y Sistema de Notificaciones", level=1)
    add_paragraph(
        document,
        "`telegram_notifier.py` cubre todo el canal Telegram. Puede enviar una alerta individual (`send_job_alert`) "
        "o un digest por lote (`send_job_alert_digest`). Tambien formatea el texto, decide a que chats mandar, "
        "trocea mensajes largos y registra resultados en la tabla `notifications`.",
    )
    document.add_heading("8.1 Formato de Mensaje", level=2)
    add_bullets(
        document,
        [
            "Cargo, empresa, portal, ubicacion y compatibilidad.",
            "Fecha publicada y fecha detectada, usando `date_utils.py` y `JOBOPS_TIMEZONE`.",
            "Motivo de match (`match_reason`).",
            "Link oficial de aplicacion.",
            "Comando sugerido para generar CV ATS.",
            "Comando sugerido para marcar el estado como `applied`.",
        ],
    )
    document.add_heading("8.2 Digest y Ruteo", level=2)
    add_paragraph(
        document,
        "El digest se ordena por score y fecha, respeta `telegram_max_message_chars`, y puede partirse en varias "
        "partes numeradas. Si el entorno define `TELEGRAM_CHAT_TARGETS`, el envio deja de ser broadcast simple "
        "y pasa a ser target-aware: cada chat puede recibir solo ofertas de ciertos roles.",
    )
    add_bullets(
        document,
        [
            "`TELEGRAM_CHAT_ID` es el fallback historico.",
            "`TELEGRAM_CHAT_IDS` permite varios chats separados por coma.",
            "`TELEGRAM_CHAT_TARGETS` usa formato `chat_id:target1,target2;chat_id2:*`.",
            "`TELEGRAM_CHAT_LABELS` usa formato `chat_id:Nombre Visible;...` para enriquecer logs y mensajes.",
        ],
    )
    document.add_heading("8.3 Cuando se Dispara", level=2)
    add_paragraph(
        document,
        "Una oferta entra a la cola de notificacion cuando: no ha sido notificada, es fresca, supera el "
        "`match_threshold` y, en LinkedIn, pasa tambien el filtro Easy Apply si esa politica esta activa. "
        "El monitor puede mandar el digest al final del ciclo completo o inmediatamente despues de cada fuente, "
        "segun `JOBOPS_NOTIFY_AFTER_EACH_SOURCE`.",
    )

    document.add_heading("9. Monitor de Ofertas Frescas", level=1)
    add_paragraph(
        document,
        "`freshness_monitor.py` es el orquestador central del scraping continuo. Toma las fuentes habilitadas, "
        "consulta cada una con el scraper correcto, decide que hacer con cada `ScrapedJob`, guarda resultados, "
        "actualiza deduplicacion y maneja notificaciones pendientes o nuevas.",
    )
    document.add_heading("9.1 Loop Operativo", level=2)
    add_mono_block(
        document,
        """monitor watch
  -> abre una sesion SQLAlchemy por ciclo
  -> get_due_sources()
  -> por cada fuente no pausada:
       -> get_scraper(portal)
       -> scraper.scrape(source)
       -> descartar por relevancia o politica LinkedIn
       -> crear/mergear oferta
       -> recalcular match
       -> registrar hash visto
       -> agregar a digest si corresponde
       -> update_source_check() o record_source_failure()
  -> enviar digest / reintentar pendientes
  -> sleep(interval)""",
    )
    document.add_heading("9.2 Manejo de Pausas y Errores", level=2)
    add_bullets(
        document,
        [
            "`CaptchaRequiredError`, `LoginRequiredError` y `SourceBlockedError` cuentan como fallos reales de fuente.",
            "Errores genericos del scraper actualizan `last_error`, pero no necesariamente pausan la fuente.",
            "Las fuentes pausadas se omiten mientras `paused_until` siga en el futuro.",
            "Si `force_all=False` y no hay fuentes due, el monitor devuelve un mensaje de ciclo vacio en lugar de fallar.",
        ],
    )
    document.add_heading("9.3 Duplicados y Pendientes", level=2)
    add_paragraph(
        document,
        "Si una oferta ya existia, el monitor no crea un nuevo registro; fusiona informacion util, agrega una nota "
        "de posible duplicado y vuelve a registrar el hash visto. Si una oferta cumple el umbral pero Telegram fallo, "
        "queda con `telegram_notified=False` y puede reenviarse con `retry_pending_alerts()` o los comandos CLI de "
        "notificaciones pendientes.",
    )

    document.add_heading("10. Configuracion (`.env`)", level=1)
    add_paragraph(
        document,
        "La configuracion real se define en `settings.py`. Algunas variables aparecen en `.env.example`; otras son "
        "soportadas por el codigo aunque no esten prelistadas en el ejemplo. Tambien hay que distinguir entre los "
        "nombres conceptuales solicitados por el usuario y los nombres reales usados por el proyecto.",
    )
    document.add_heading("10.1 Variables Soportadas por `settings.py`", level=2)
    add_table(
        document,
        ["Variable", "Efecto real en el codigo"],
        [
            ("JOBOPS_DB_PATH", "Ruta del archivo SQLite principal."),
            ("JOBOPS_MATCH_THRESHOLD", "Score minimo para considerar una oferta notificable."),
            ("JOBOPS_SCRAPER_TIMEOUT", "Timeout HTTP para scrapers requests."),
            ("JOBOPS_SCRAPER_USER_AGENT", "User-Agent para requests y tambien para Chrome Selenium."),
            ("JOBOPS_MAX_RESULTS_PER_SOURCE", "Maximo de ofertas por fuente por corrida."),
            ("JOBOPS_MIN_MONITOR_INTERVAL_MINUTES", "Intervalo minimo permitido al registrar o actualizar fuentes."),
            ("JOBOPS_TIMEZONE", "Zona horaria usada en formatos de fecha/hora."),
            ("JOBOPS_TELEGRAM_DIGEST_LIMIT", "Limite preferido de ofertas por digest; tiene precedencia sobre la variable legacy."),
            ("JOBOPS_TELEGRAM_DIGEST_MAX_JOBS", "Fallback legacy para el limite de digest si la variable anterior no existe."),
            ("JOBOPS_TELEGRAM_MAX_MESSAGE_CHARS", "Corte aproximado de caracteres antes de dividir un digest."),
            ("JOBOPS_NOTIFY_AFTER_EACH_SOURCE", "Si es true, manda digest inmediato por fuente y luego procesa el remanente."),
            ("JOBOPS_ENABLE_SELENIUM", "Habilita scrapers Selenium."),
            ("JOBOPS_SELENIUM_HEADLESS", "Controla si Chrome se abre sin UI."),
            ("JOBOPS_SELENIUM_PAGE_LOAD_TIMEOUT", "Timeout de carga de pagina para Selenium."),
            ("JOBOPS_SELENIUM_SCROLL_PAUSE", "Pausa entre scrolls o waits fallback."),
            ("JOBOPS_SELENIUM_MAX_SCROLLS", "Cantidad maxima de scrolls automaticos."),
            ("JOBOPS_SELENIUM_USER_DATA_DIR", "User data dir de Chrome para reutilizar cookies/sesion local."),
            ("JOBOPS_SELENIUM_PROFILE_DIRECTORY", "Subperfil de Chrome dentro del user data dir."),
            ("JOBOPS_LINKEDIN_FETCH_DETAILS", "Si es true, abre cada oferta LinkedIn para leer la descripcion de detalle."),
            ("JOBOPS_LINKEDIN_ONLY_EASY_APPLY", "Si es true, descarta ofertas LinkedIn que no sean Easy Apply."),
            ("TELEGRAM_BOT_TOKEN", "Token del bot de Telegram."),
            ("TELEGRAM_CHAT_ID", "Chat unico historico."),
            ("TELEGRAM_CHAT_IDS", "Lista de chats separada por comas."),
            ("TELEGRAM_CHAT_TARGETS", "Mapa chat -> targets, separado por `;` y `,`."),
            ("TELEGRAM_CHAT_LABELS", "Mapa chat -> etiqueta legible."),
            ("GMAIL_EMAIL", "Cuenta usada para el workflow diario."),
            ("GMAIL_APP_PASSWORD", "App password asociada a Gmail."),
        ],
    )
    document.add_heading("10.2 Traduccion de Nombres Conceptuales", level=2)
    add_bullets(
        document,
        [
            "El nombre real de `DB_PATH` en el codigo es `JOBOPS_DB_PATH`.",
            "El nombre real de `ENABLE_SELENIUM` es `JOBOPS_ENABLE_SELENIUM`.",
            "El nombre real de `HEADLESS` es `JOBOPS_SELENIUM_HEADLESS`.",
            "No existe una variable literal `PROFILE_DIR`; el codigo usa `JOBOPS_SELENIUM_USER_DATA_DIR` y `JOBOPS_SELENIUM_PROFILE_DIRECTORY`.",
        ],
    )
    document.add_heading("10.3 Nota sobre `.env.example`", level=2)
    add_paragraph(
        document,
        "`.env.example` documenta la configuracion principal, pero no lista todos los campos que `settings.py` ya soporta. "
        "Por ejemplo, el codigo reconoce `JOBOPS_SELENIUM_USER_DATA_DIR`, `JOBOPS_SELENIUM_PROFILE_DIRECTORY` y "
        "`JOBOPS_LINKEDIN_FETCH_DETAILS`, aunque no aparezcan en ese archivo de ejemplo.",
    )

    document.add_heading("11. Casos Reales del Sistema", level=1)
    document.add_heading("11.1 LinkedIn devuelve 0 jobs", level=2)
    add_paragraph(
        document,
        "En el estado actual del codigo, 0 jobs en LinkedIn no significa automaticamente bloqueo. "
        "Si `linkedin_selenium` detecta texto real de pagina vacia (`no results found`, `no matching jobs found`, etc.), "
        "no lanza excepcion: simplemente devuelve cero resultados y la fuente sigue activa. "
        "En cambio, si no hay cards y aparecen senales de authwall, login o captcha, el scraper lanza una excepcion "
        "de bloqueo/login y esa fuente si aumenta `failure_count`.",
    )
    document.add_heading("11.2 Hay captcha o access wall", level=2)
    add_paragraph(
        document,
        "En scrapers requests, `403` y `429` pasan por `BaseJobScraper._request_text()` y se convierten en bloqueo real. "
        "En LinkedIn Selenium y Computrabajo Selenium, la decision depende del HTML renderizado: si aparecen senales "
        "de captcha, checkpoint, login wall o access denied, se lanzan excepciones especificas. El monitor interpreta "
        "esas excepciones como fallos reales de fuente y puede pausar la fuente al tercer evento consecutivo.",
    )
    document.add_heading("11.3 Se descartan ofertas", level=2)
    add_paragraph(
        document,
        "Cuando una oferta no coincide con el target, no desaparece silenciosamente. El sistema crea un registro en "
        "`discarded_jobs` con razones, keywords y score preliminar. Eso permite revisar despues si el matcher estaba "
        "siendo demasiado estricto. Tambien se pueden exportar descartadas o reprocesarlas con reglas actualizadas.",
    )
    document.add_heading("11.4 Se envia Telegram", level=2)
    add_paragraph(
        document,
        "Cuando una oferta valida supera el umbral, el monitor la pone en una cola de digest. Si Telegram entrega el "
        "mensaje, `job_service.mark_offer_telegram_notified()` marca la oferta como notificada y "
        "`telegram_notifier.register_notification()` deja un historial en `notifications`. Si el envio falla para todos "
        "los chats, la oferta queda pendiente y puede reintentarse sin volver a scrapearla.",
    )
    document.add_heading("11.5 Computrabajo no encuentra cards visibles", level=2)
    add_paragraph(
        document,
        "Computrabajo ya no trata `len(job_cards) == 0` como bloqueo por defecto. El scraper primero revisa links "
        "relevantes, HTML renderizado, patrones de captcha/login y senales de sin resultados. Solo si hay evidencia "
        "real marca `blocked`; si la pagina parece valida pero cambio el DOM, devuelve `scraper_broken`; y si la "
        "pagina indica explicitamente vacio, devuelve `no_results`.",
    )
    document.add_heading("11.6 Gmail diario", level=2)
    add_paragraph(
        document,
        "El flujo `scan-daily` existe y `workflows.py` puede convertir alertas en ofertas y disparar Telegram. "
        "Sin embargo, `gmail_reader.py` hoy devuelve una lista vacia de manera conservadora incluso si existen "
        "credenciales. Arquitectonicamente el pipeline esta preparado; funcionalmente la extraccion IMAP real aun "
        "no esta implementada.",
    )

    document.add_heading("12. Conclusion", level=1)
    add_paragraph(
        document,
        "El repositorio actual muestra un sistema CLI maduro para uso personal/local, con persistencia SQLite, "
        "monitoreo incremental, deduplicacion, auditoria de descartes, notificacion por Telegram y una cadena "
        "separada para importacion de hoja de vida y generacion de CV ATS. La arquitectura es simple de desplegar "
        "porque no depende de servicios remotos propios ni de colas distribuidas; a cambio, concentra mucha "
        "responsabilidad en el proceso local y en la sesion SQLite.",
    )
    add_paragraph(
        document,
        "Los componentes mas avanzados son hoy el monitor, el matcher, la auditoria de descartes, y los scrapers "
        "Selenium de LinkedIn, Indeed y Computrabajo. El componente mas claramente incompleto es Gmail: existe el "
        "workflow, pero no la lectura real de correos. En estructura fisica, el proyecto sigue una organizacion plana "
        "y pragmatica, suficiente para un MVP extendido, aunque ya hay volumen suficiente como para justificar una "
        "separacion futura en subpackages por dominio.",
    )
    document.add_heading("12.1 Mejoras Futuras Compatibles con el Estado Actual", level=2)
    add_bullets(
        document,
        [
            "Implementar lectura IMAP real en `gmail_reader.py`.",
            "Agregar dashboard web o vista de administracion sobre SQLite.",
            "Agregar reportes periodicos y metricas de fuentes.",
            "Separar el package en submodulos por dominio (`monitor`, `notifications`, `resume`, `scrapers`).",
            "Agregar Docker y CI, alineado con el roadmap ya documentado en README.",
            "Fortalecer observabilidad con logging estructurado y metricas por fuente.",
        ],
    )
    document.add_heading("12.2 Mejoras que Requeririan Cambiar la Politica del Proyecto", level=2)
    add_paragraph(
        document,
        "La peticion menciona ideas como anti-captcha, rotacion de perfiles o proxies. Esas capacidades no existen "
        "en el codigo actual y ademas chocan con la postura de scraping responsable declarada por el propio proyecto. "
        "Si alguna vez se evaluaran, tendrian que pasar por una decision explicita de producto y de cumplimiento; no "
        "son una evolucion natural del estado actual del repositorio.",
    )

    document.add_heading("Apendice A. Matriz de Portales Soportados", level=1)
    add_table(
        document,
        ["Portal", "Tecnica", "Observacion"],
        [
            ("linkedin", "Requests + selectores", "HTML publico sencillo; menor tolerancia al bloqueo."),
            ("linkedin_selenium", "Selenium", "Cards publicas o layout visible con perfil local de Chrome; deteccion de authwall y tipo de solicitud."),
            ("indeed", "Requests + selectores", "Lectura publica basica."),
            ("indeed_selenium", "Selenium", "Extraccion de `jk` y URL canonica `viewjob`."),
            ("computrabajo", "Selenium", "HTML renderizado, extraccion robusta y detalle enriquecido."),
            ("elempleo", "Requests + selectores", "Selector-based con heuristicas adicionales."),
            ("magneto", "Requests + selectores", "Selector-based simple."),
            ("torre", "Requests + selectores", "Selector-based simple."),
            ("getonboard", "Requests + selectores", "Selector-based simple."),
            ("sena", "Requests + selectores", "Selector-based simple."),
        ],
    )

    document.add_heading("Apendice B. Suite de Pruebas Visible en el Repositorio", level=1)
    add_paragraph(
        document,
        "La carpeta `tests/` contiene pruebas especificas para matcher, monitor, servicios de ofertas, descarte, "
        "scrapers requests, scrapers Selenium, Telegram y el pipeline de CV/resume. No es una garantia de cobertura "
        "total, pero si indica que el proyecto ya valida formalmente sus rutas mas importantes.",
    )
    add_bullets(
        document,
        [
            "`test_freshness_monitor.py`.",
            "`test_search_sources.py`.",
            "`test_scrapers.py` y `test_selenium_scrapers.py`.",
            "`test_matcher.py`, `test_discarded_jobs.py`, `test_job_service.py`, `test_offer_management.py`.",
            "`test_telegram_notifier.py`.",
            "`test_resume_reader.py` y `test_ats_resume_builder.py`.",
        ],
    )

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
