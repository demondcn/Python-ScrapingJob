from pathlib import Path

import pytest
from docx import Document

from src.jobops_assistant.ats_resume_builder import build_ats_filename, build_ats_resume, get_available_targets, rank_skills, select_relevant_experiences
from src.jobops_assistant.models import JobOffer
from src.jobops_assistant.resume_profile_service import load_resume_profile
from src.jobops_assistant.schemas import ResumeExperience, ResumeProfile, ResumeProject


def _sample_profile() -> ResumeProfile:
    return ResumeProfile(
        full_name="Cristian Stiven Guerrero Andrade",
        location="Soacha, Colombia",
        phone="+57 305 423 3742",
        email="guerrero70407@gmail.com",
        linkedin="https://linkedin.com/in/demondcn/",
        github="https://github.com/demondcn",
        portfolio="https://emprex360.vercel.app",
        professional_summary="Tecnólogo en Desarrollo de Software.",
        technical_skills=[
            "Python",
            "Java",
            "JavaScript",
            "TypeScript",
            "SQL",
            "PostgreSQL",
            "MySQL",
            "Git",
            "GitHub",
            "Linux",
            "Docker",
            "React",
            "Next.js",
            "Vercel",
            "Neon",
            "R",
            "de",
            "en",
            "Software",
            "Web",
        ],
        soft_skills=["Documentación técnica", "Resolución de problemas", "Trabajo en equipo"],
        experiences=[
            ResumeExperience(
                role="Auxiliar de Oficina II",
                company="Universidad de Cundinamarca",
                location="Soacha, Colombia",
                start_date="Febrero 2025",
                end_date="Junio 2025",
                bullets=[
                    "Registro y control de inventarios, reportes y novedades del centro de cómputo.",
                    "Soporte básico a usuarios en el uso de herramientas informáticas y equipos.",
                ],
            ),
            ResumeExperience(
                role="Auxiliar de Oficina II",
                company="Universidad de Cundinamarca",
                location="Soacha, Colombia",
                start_date="Agosto 2025",
                end_date="Noviembre 2025",
                bullets=[
                    "Apoyo operativo y soporte administrativo en los centros de cómputo.",
                    "Gestión de recursos tecnológicos, inventario y reportes.",
                    "Automatización de procesos de apagado de equipos.",
                    "Soporte básico a usuarios y diagnóstico de incidencias.",
                ],
            ),
            ResumeExperience(
                role="Técnico en Soporte de Equipos de Cómputo",
                company="Universidad de Cundinamarca",
                location="Soacha, Colombia",
                start_date="Marzo 2023",
                end_date="Junio 2023",
                bullets=[
                    "Mantenimiento preventivo y correctivo de hardware y software.",
                    "Documentación de configuraciones y hojas de vida de equipos.",
                    "Resolución de fallos y soporte técnico a usuarios.",
                ],
            ),
            ResumeExperience(
                role="Desarrollador Full Stack Web",
                company="Emprex360",
                location="Remoto",
                start_date="Enero 2024",
                end_date="Febrero 2025",
                bullets=[
                    "Implementación de dashboards interactivos con React y Next.js.",
                    "Despliegue en Vercel con base de datos alojada en Neon.",
                    "Carga de datos en tiempo real con PostgreSQL y MySQL.",
                    "Documentación técnica y resolución de problemas.",
                ],
            ),
            ResumeExperience(
                role="Desarrollador Frontend y Mobile",
                company="Chocontano Restaurante",
                location="Soacha, Colombia",
                start_date="Marzo 2024",
                end_date="Junio 2024",
                bullets=[
                    "Desarrollo de aplicación móvil con React Native y sistema E-Commerce web.",
                    "Implementación de pedidos en línea, inventario y panel administrativo.",
                ],
            ),
            ResumeExperience(
                role="Desarrollador Android",
                company="Charles Barber",
                location="Bogotá, Colombia",
                start_date="Agosto 2024",
                end_date="Octubre 2024",
                bullets=[
                    "Desarrollo de aplicación Android en Java para control de inventario y ventas.",
                    "Implementación de módulo de tienda en línea y gestión administrativa.",
                ],
            ),
            ResumeExperience(
                role="Desarrollador Frontend Web",
                company="Kepri Holística",
                location="Soacha, Colombia",
                start_date="Octubre 2024",
                end_date="Noviembre 2024",
                bullets=[
                    "Participación en bootcamp UDEC para desarrollar una página web E-Commerce con Next.js.",
                    "La plataforma incluía catálogo de servicios, información institucional y sistema de pagos.",
                    "Diseño responsivo orientado a experiencia de usuario.",
                ],
            ),
            ResumeExperience(
                role="Operador",
                company="Salitre Magico",
                location="Bogotá, Colombia",
                start_date="Enero 2024",
                end_date="Julio 2024",
                bullets=[
                    "Controlar el flujo de entrada y salida de los vehículos en el estacionamiento.",
                    "Brindar información real y verídica a los visitantes del parque.",
                ],
            ),
        ],
        projects=[
            ResumeProject(
                name="Emprex360",
                role="Full Stack",
                technologies=["React", "Next.js", "PostgreSQL", "Vercel", "Neon"],
                bullets=["Plataforma web para análisis de datos y dashboards."],
                url="https://emprex360.vercel.app",
            )
        ],
        education=["Universidad de Cundinamarca - Tecnólogo en Desarrollo de Software"],
        certifications=["Curso Profesional de Git y GitHub", "Curso de Introducción a la Terminal y Línea de Comandos"],
        languages=["Español - Nativo", "Inglés - Lectura técnica"],
        raw_text=(
            "JavaScript TypeScript React Next.js Python SQL PostgreSQL MySQL Git GitHub Linux Docker "
            "Vercel Neon soporte técnico soporte a usuarios documentación técnica línea de comandos "
            "inventario mantenimiento preventivo mantenimiento correctivo automatización incidencias "
            "resolución de problemas reportes centros de cómputo cloud despliegue aplicaciones web "
            "Más información Daniel Pinto 3219518649 Agudelo Davalos Jesus Pavon Jesus Blanco "
            "Salitre Magico vehículos parque Referencias raw_text"
        ),
    )


def _banned_fragments() -> list[str]:
    return [
        "Con fortalezas en R",
        "R, en, de",
        "R, GitHub, Git, de",
        "R, de, en",
        "Tecnologo",
        "tecnico",
        "documentacion",
        "Educacion",
        "Más información",
        "Mas informacion",
        "Referencias",
        "Daniel Pinto",
        "Agudelo",
        "Jesus Pavon",
        "Jesus Blanco",
        "Salitre Magico",
        "vehículos",
        "vehiculos",
        "parque",
        "raw_text",
    ]


def test_build_ats_resume_creates_all_targets_without_noise(tmp_path: Path):
    profile = _sample_profile()

    for target in get_available_targets():
        output = tmp_path / f"{target}.docx"
        build_ats_resume(profile, target, output)

        assert output.exists()
        document = Document(output)
        joined = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert "Habilidades Técnicas" in joined
        assert "Experiencia Profesional y Proyectos" in joined
        assert "Experiencia en Soporte e Infraestructura" not in joined
        assert "●" not in joined
        assert len(document.tables) == 0
        assert len(document.inline_shapes) == 0
        for fragment in _banned_fragments():
            assert fragment not in joined


def test_build_ats_resume_with_job_offer_uses_clean_builder(tmp_path: Path):
    output = tmp_path / "job.docx"
    offer = JobOffer(
        id=1,
        title="Soporte de Aplicaciones Junior",
        company="ABC",
        portal="Computrabajo",
        location="Bogotá",
        modality="Híbrido",
        salary="",
        url="https://example.com/offer",
        description="Soporte a usuarios, tickets, SQL, documentación e incidentes",
        requirements="Junior",
    )

    build_ats_resume(_sample_profile(), "soporte_aplicaciones", output, offer)
    joined = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)

    assert "Adaptado para una vacante de Soporte de Aplicaciones Junior." in joined
    assert "Soporte y aplicaciones:" in joined
    assert "Bases de datos:" in joined
    assert "Herramientas:" in joined
    for fragment in _banned_fragments():
        assert fragment not in joined


def test_select_relevant_experiences_orders_support_targets():
    offer = JobOffer(
        id=1,
        title="Soporte de Aplicaciones Junior",
        company="ABC",
        portal="Computrabajo",
        location="Bogotá",
        modality="Híbrido",
        salary="",
        url="https://example.com/offer",
        description="Soporte a usuarios, tickets, SQL, documentación e incidentes",
        requirements="Junior",
    )
    expected = [
        ("Universidad de Cundinamarca", "Auxiliar de Oficina II"),
        ("Universidad de Cundinamarca", "Técnico en Soporte de Equipos de Cómputo"),
        ("Emprex360", "Desarrollador Full Stack Web"),
        ("Chocontano Restaurante", "Desarrollador Frontend y Mobile"),
        ("Charles Barber", "Desarrollador Android"),
        ("Kepri Holística", "Desarrollador Frontend Web"),
    ]

    support = select_relevant_experiences(_sample_profile(), "soporte_aplicaciones")
    support_job = select_relevant_experiences(_sample_profile(), "soporte_aplicaciones", offer)

    assert [(item.company, item.role) for item in support] == expected
    assert [(item.company, item.role) for item in support_job] == expected


def test_select_relevant_experiences_orders_infraestructura():
    expected = [
        ("Universidad de Cundinamarca", "Técnico en Soporte de Equipos de Cómputo"),
        ("Universidad de Cundinamarca", "Auxiliar de Oficina II"),
        ("Emprex360", "Desarrollador Full Stack Web"),
        ("Chocontano Restaurante", "Desarrollador Frontend y Mobile"),
        ("Charles Barber", "Desarrollador Android"),
        ("Kepri Holística", "Desarrollador Frontend Web"),
    ]

    selected = select_relevant_experiences(_sample_profile(), "infraestructura_junior")

    assert [(item.company, item.role) for item in selected] == expected


def test_select_relevant_experiences_orders_devops():
    expected = [
        ("Emprex360", "Desarrollador Full Stack Web"),
        ("Kepri Holística", "Desarrollador Frontend Web"),
        ("Universidad de Cundinamarca", "Auxiliar de Oficina II"),
        ("Chocontano Restaurante", "Desarrollador Frontend y Mobile"),
        ("Charles Barber", "Desarrollador Android"),
    ]

    selected = select_relevant_experiences(_sample_profile(), "devops_trainee")

    assert [(item.company, item.role) for item in selected] == expected


def test_select_relevant_experiences_use_canonical_clean_bullets():
    selected = select_relevant_experiences(_sample_profile(), "soporte_aplicaciones")
    experience_map = {(item.company, item.role): item.bullets for item in selected}

    assert experience_map[("Emprex360", "Desarrollador Full Stack Web")] == [
        "Desarrollo de plataforma web para análisis de datos y diagnóstico empresarial.",
        "Implementación de dashboards interactivos y carga de datos en tiempo real con React, Next.js y PostgreSQL.",
        "Despliegue en Vercel con base de datos alojada en Neon.",
    ]
    assert experience_map[("Chocontano Restaurante", "Desarrollador Frontend y Mobile")] == [
        "Desarrollo de aplicación móvil con React Native y sistema E-Commerce web.",
        "Implementación de pedidos en línea, inventario y panel administrativo.",
    ]
    assert experience_map[("Charles Barber", "Desarrollador Android")] == [
        "Desarrollo de aplicación Android en Java para control de inventario y ventas.",
        "Implementación de módulo de tienda en línea y gestión administrativa.",
    ]
    assert experience_map[("Universidad de Cundinamarca", "Auxiliar de Oficina II")] == [
        "Apoyo operativo y soporte a usuarios en centros de cómputo.",
        "Gestión e inventario de recursos tecnológicos en laboratorios.",
        "Automatización de procesos de apagado de equipos para optimizar la operación.",
    ]
    assert experience_map[("Universidad de Cundinamarca", "Técnico en Soporte de Equipos de Cómputo")] == [
        "Mantenimiento preventivo y correctivo de equipos de cómputo.",
        "Documentación de historial de mantenimiento, actualizaciones y configuraciones.",
        "Apoyo en diagnóstico de fallos de hardware y software.",
    ]
    assert experience_map[("Kepri Holística", "Desarrollador Frontend Web")] == [
        "Desarrollo de página web E-Commerce con Next.js.",
        "Implementación de catálogo de servicios, información institucional y diseño responsivo.",
    ]
    assert all(len(item.bullets) <= 3 for item in selected)


def test_rank_skills_filters_noise_and_duplicates():
    skills = rank_skills(_sample_profile(), "devops_trainee")
    assert "R" not in skills
    assert "de" not in skills
    assert "en" not in skills
    assert "Software" not in skills
    assert len(skills) == len(set(skills))


def test_build_ats_filename_with_job_offer():
    offer = JobOffer(id=1, title="Soporte de Aplicaciones Junior", company="ABC", portal="", location="", modality="", salary="", url="https://example.com")
    file_name = build_ats_filename(_sample_profile(), "soporte_aplicaciones", offer)
    assert file_name == "CV_ATS_Cristian_Stiven_Guerrero_Andrade_Oferta_1_Soporte_de_Aplicaciones_Junior.docx"


def test_load_resume_profile_raises_if_missing(tmp_path: Path):
    with pytest.raises(FileNotFoundError):
        load_resume_profile(tmp_path / "missing.json")
