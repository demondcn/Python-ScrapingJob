from pathlib import Path

from docx import Document

from src.jobops_assistant.resume_reader import clean_resume_text, read_resume_file


def test_clean_resume_text_removes_extra_spacing():
    raw = "Cristian   Guerrero\r\n\r\nEmail: test@example.com\r\n\r\n\r\nLinkedIn: linkedin.com/in/demo"
    cleaned = clean_resume_text(raw)
    assert "  " not in cleaned
    assert "\n\n\n" not in cleaned


def test_read_docx_detects_contact_and_excludes_references(tmp_path: Path):
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Cristian Stiven Guerrero Andrade")
    document.add_paragraph("Soacha, Colombia | Tel: +57 305 423 3742 | Email: guerrero70407@gmail.com | LinkedIn: linkedin.com/in/demondcn/")
    document.add_paragraph("EDUCACION")
    document.add_paragraph("Universidad de Cundinamarca - Tecnologo en Desarrollo de Software")
    document.add_paragraph("EXPERIENCIA PROFESIONAL")
    document.add_paragraph("Emprex360 - Desarrollador Full Stack Web")
    document.add_paragraph("Soacha, Colombia | Enero 2024 - Febrero 2025")
    document.add_paragraph("• Desarrollo de dashboards y despliegue en Vercel.")
    document.add_paragraph("HABILIDADES Y COMPETENCIAS")
    document.add_paragraph("Lenguajes y Tecnologias:")
    document.add_paragraph("Python, SQL, Git, Linux, Docker")
    document.add_paragraph("Habilidades Blandas:")
    document.add_paragraph("Comunicacion efectiva, trabajo en equipo")
    document.add_paragraph("IDIOMAS")
    document.add_paragraph("Español - Nativo | Ingles - Lectura tecnica")
    document.add_paragraph("MÁS INFORMACIÓN")
    document.add_paragraph("Daniel Pinto - 3219518649")
    document.add_paragraph("DISPONIBILIDAD")
    document.add_paragraph("Inmediata")
    document.add_paragraph("REFERENCIAS")
    document.add_paragraph("Referencia familiar: Maria Perez")
    document.add_paragraph("Operador - Salitre Magico")
    document.add_paragraph("Bogotá, Colombia | Enero 2024 - Julio 2024")
    document.add_paragraph("• Elaboración de planillas para monitorear el estado de los vehículos del parque.")
    document.save(path)

    profile = read_resume_file(path)

    assert profile.full_name == "Cristian Stiven Guerrero Andrade"
    assert profile.email == "guerrero70407@gmail.com"
    assert "305 423 3742" in profile.phone
    assert "linkedin.com/in/demondcn/" in profile.linkedin
    assert profile.experiences
    assert all("Referencia" not in item for item in profile.certifications)
    assert all("Más información" not in item for item in profile.languages)
    assert all("Daniel Pinto" not in item for item in profile.languages)
    assert all("Salitre" not in f"{experience.role} {experience.company}" for experience in profile.experiences)
    assert "Python" in profile.technical_skills


def test_read_docx_normalizes_known_experiences_and_merges_auxiliar(tmp_path: Path):
    path = tmp_path / "resume_experiences.docx"
    document = Document()
    document.add_paragraph("Cristian Stiven Guerrero Andrade")
    document.add_paragraph("Soacha, Colombia | Tel: +57 305 423 3742 | Email: guerrero70407@gmail.com")
    document.add_paragraph("EXPERIENCIA PROFESIONAL")
    document.add_paragraph("Kepri Holística - Bootcamp UDEC Desarrollador FrontEnd Web Soacha, Colombia")
    document.add_paragraph("Remoto Febrero 2024")
    document.add_paragraph("• Participé en un bootcamp liderado por la Universidad de Cundinamarca.")
    document.add_paragraph("• E-commerce con Next.js para un centro de restauración espiritual.")
    document.add_paragraph("• La plataforma incluía catálogo de servicios y sistema de pagos.")
    document.add_paragraph("• Diseñé y desarrollé una aplicación móvil para el estudio del lactato en sangre.")
    document.add_paragraph("Charles Barber - App Android Desarrollador Android y E-Commerce Bogotá, Colombia")
    document.add_paragraph("Java, Android Studio | Agosto 2024 - Octubre 2024")
    document.add_paragraph("• Diseñé y desarrollé una aplicación móvil especializada para barberías.")
    document.add_paragraph("• Incluí módulo de tienda en línea y gestión de productos desde el lado administrativo.")
    document.add_paragraph("Universidad de Cundinamarca - Auxiliar de Oficina II")
    document.add_paragraph("Soacha, Colombia | Agosto 2025 - Noviembre 2025")
    document.add_paragraph("• Apoyo operativo y soporte administrativo en los centros de cómputo.")
    document.add_paragraph("Universidad de Cundinamarca - Auxiliar de Oficina II")
    document.add_paragraph("Soacha, Colombia | Febrero 2025 - Junio 2025")
    document.add_paragraph("• Registro y control de inventarios, reportes y novedades del centro de cómputo.")
    document.add_paragraph("Universidad de Cundinamarca Fusagasugá - Desarrollador Full Stack")
    document.add_paragraph("Enero 2024 - Febrero 2025")
    document.add_paragraph("• Proyecto: Emprex360 diagnóstico para Empresas.")
    document.save(path)

    profile = read_resume_file(path)

    experiencias = {(item.company, item.role): item for item in profile.experiences}
    assert ("Kepri Holística", "Desarrollador Frontend Web") in experiencias
    assert experiencias[("Kepri Holística", "Desarrollador Frontend Web")].location == "Soacha, Colombia"
    assert experiencias[("Kepri Holística", "Desarrollador Frontend Web")].start_date == "Octubre 2024"
    assert experiencias[("Kepri Holística", "Desarrollador Frontend Web")].end_date == "Noviembre 2024"
    assert all("lactato" not in bullet.lower() for bullet in experiencias[("Kepri Holística", "Desarrollador Frontend Web")].bullets)

    assert ("Charles Barber", "Desarrollador Android") in experiencias
    assert experiencias[("Charles Barber", "Desarrollador Android")].location == "Bogotá, Colombia"

    assert ("Universidad de Cundinamarca", "Auxiliar de Oficina II") in experiencias
    auxiliar = experiencias[("Universidad de Cundinamarca", "Auxiliar de Oficina II")]
    assert auxiliar.start_date == "Febrero 2025"
    assert auxiliar.end_date == "Noviembre 2025"
    assert len(auxiliar.bullets) == 2

    assert ("Emprex360", "Desarrollador Full Stack Web") in experiencias
